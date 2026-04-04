from __future__ import annotations

import re
from io import BytesIO
from pathlib import Path
from urllib.parse import urlparse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from xml.sax.saxutils import escape
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
import requests

from app.config import settings


FONT_CANDIDATES = (
    r"C:\Windows\Fonts\arial.ttf",
    "/usr/share/fonts/truetype/liberation2/LiberationSans-Regular.ttf",
    "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
    "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
)

_ALLOWED_LOGO_HOSTS = {
    host.strip().lower()
    for host in settings.document_logo_allowed_hosts.split(",")
    if host.strip()
}


def _is_allowed_logo_url(url: str) -> bool:
    """Дозволяє лише HTTPS URL з явного allowlist хостів."""
    try:
        parsed = urlparse(url)
        hostname = (parsed.hostname or "").strip().lower()
        return parsed.scheme == "https" and bool(hostname) and hostname in _ALLOWED_LOGO_HOSTS
    except Exception:
        return False


def sanitize_filename(value: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._-]+", "_", value.strip())
    cleaned = cleaned.strip("._")
    return cleaned or "document"


def _resolve_font() -> str:
    for idx, font_path in enumerate(FONT_CANDIDATES, start=1):
        if not Path(font_path).exists():
            continue
        font_name = f"LegalAIFont{idx}"
        try:
            if font_name not in pdfmetrics.getRegisteredFontNames():
                pdfmetrics.registerFont(TTFont(font_name, font_path))
            return font_name
        except Exception:
            continue
    return "Helvetica"


def render_docx_bytes(*, title: str, text: str, logo_url: str | None = None) -> bytes:
    doc = Document()
    
    # logo_url embedding (лише для HTTPS host allowlist)
    if logo_url and _is_allowed_logo_url(logo_url):
        try:
            resp = requests.get(logo_url, timeout=5)
            if resp.status_code == 200:
                img_stream = BytesIO(resp.content)
                header = doc.sections[0].header
                p = header.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                r = p.add_run()
                r.add_picture(img_stream, width=Inches(1.5))
        except Exception:
            pass # Fail silently on invalid logos
    
    # Branded Header
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(0.75)

    # Professional Title (Header style)
    t = doc.add_heading(title or "Юридичний документ", level=1)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Body
    for para in text.splitlines():
        para = para.strip()
        if not para:
            doc.add_paragraph()
            continue
            
        p = doc.add_paragraph()
        p.style.font.name = "Arial"
        p.style.font.size = Pt(12)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.first_line_indent = Inches(0.5)

        # Basic Markdown Bold handling
        parts = re.split(r'(\*\*.*?\*\*)', para)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                p.add_run(part)

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


def render_pdf_bytes(*, title: str, text: str, logo_url: str | None = None) -> bytes:
    buffer = BytesIO()
    font_name = _resolve_font()
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=3*cm,
        topMargin=2*cm,
        bottomMargin=2*cm
    )
    
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'LegalTitle',
        parent=styles['Heading1'],
        fontName=font_name,
        fontSize=14,
        alignment=1, # Center
        spaceAfter=20,
        leading=18
    )
    
    body_style = ParagraphStyle(
        'LegalBody',
        parent=styles['Normal'],
        fontName=font_name,
        fontSize=12,
        alignment=4, # Justify
        leading=16,
        firstLineIndent=1.27*cm,
        spaceAfter=10
    )
    
    content = []
    
    # Logo embedding (лише для HTTPS host allowlist)
    if logo_url and _is_allowed_logo_url(logo_url):
        try:
            resp = requests.get(logo_url, timeout=5)
            if resp.status_code == 200:
                img_stream = BytesIO(resp.content)
                content.append(Image(img_stream, width=4*cm, height=2*cm, hAlign='RIGHT'))
                content.append(Spacer(1, 10))
        except Exception:
            pass
            
    content.append(Paragraph(escape(title or "Юридичний документ"), title_style))
    
    for para in text.splitlines():
        para = para.strip()
        if not para:
            content.append(Spacer(1, 10))
            continue
            
        # Convert **text** to <b>text</b> for reportlab
        styled_para = escape(para)
        styled_para = styled_para.replace("&lt;b&gt;", "<b>").replace("&lt;/b&gt;", "</b>") # Undo escape for manual tags if any
        styled_para = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', styled_para)
        
        content.append(Paragraph(styled_para, body_style))
        
    doc.build(content)
    return buffer.getvalue()
