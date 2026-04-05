"""
Legal AI Platform — self-contained FastAPI backend.

All routes are defined here so the build context is just backend/
(no frontend/ copy needed).
"""
from __future__ import annotations

import base64
from collections import defaultdict, deque
import hashlib
import hmac
import json
import os
import re
import uuid
from datetime import datetime, timedelta, timezone
from typing import Any, Optional

import asyncio
import pathlib
import time

import httpx
from fastapi import Depends, FastAPI, File, Form, HTTPException, Query, Request, UploadFile, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.security import HTTPAuthorizationCredentials, HTTPBearer
from pydantic import BaseModel
from sqlalchemy import text
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.v1.router import api_v1_router
from app.core.celery_app import celery_app
from app.db import Base, engine, get_session
from app.tasks.ai_jobs import analyze_intake_job, generate_document_job

# ── Upload directory (ephemeral but safe for single request lifecycle) ─────────
UPLOAD_DIR = pathlib.Path(os.getenv("UPLOAD_DIR", "/tmp/uploads"))
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# ── App ───────────────────────────────────────────────────────────────────────
app = FastAPI(title="Legal AI Platform", version="1.0.0")
app.include_router(api_v1_router, prefix="/api/v1")

FRONTEND_URL = os.getenv("FRONTEND_URL", "https://frontend-production-459a.up.railway.app")
ALLOWED_ORIGINS = [
    FRONTEND_URL,
    "http://localhost:3000",
    "http://127.0.0.1:3000",
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=ALLOWED_ORIGINS,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# ── Startup: create tables ────────────────────────────────────────────────────
@app.on_event("startup")
async def startup():
    try:
        async with engine.begin() as conn:
            # Import models so Base knows about them
            from app.models.user import User  # noqa: F401
            from app.models.case import Case  # noqa: F401
            await conn.run_sync(Base.metadata.create_all)
            # Temporary backward-compat fallback for older environments.
            # New schema changes MUST go through Alembic revisions.
            if os.getenv("STARTUP_DDL_FALLBACK", "true").lower() == "true":
                for stmt in _MIGRATIONS:
                    try:
                        await conn.execute(text(stmt))
                    except Exception:
                        pass
    except Exception as e:
        print(f"[startup] DB init skipped: {e}")

_MIGRATIONS = [
    "ALTER TABLE users ADD COLUMN IF NOT EXISTS password_hash TEXT",
    "ALTER TABLE users ADD COLUMN IF NOT EXISTS updated_at TIMESTAMPTZ DEFAULT NOW()",
    "ALTER TABLE users ADD COLUMN IF NOT EXISTS company TEXT",
    "ALTER TABLE users ADD COLUMN IF NOT EXISTS entity_type TEXT",
    "ALTER TABLE users ADD COLUMN IF NOT EXISTS tax_id TEXT",
    "ALTER TABLE users ADD COLUMN IF NOT EXISTS address TEXT",
    "ALTER TABLE users ADD COLUMN IF NOT EXISTS phone TEXT",
    "ALTER TABLE users ADD COLUMN IF NOT EXISTS logo_url TEXT",
    """CREATE TABLE IF NOT EXISTS subscriptions (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        plan TEXT NOT NULL DEFAULT 'FREE',
        docs_used INT NOT NULL DEFAULT 0,
        docs_limit INT,
        period_start TIMESTAMPTZ DEFAULT NOW(),
        period_end TIMESTAMPTZ,
        created_at TIMESTAMPTZ DEFAULT NOW(),
        updated_at TIMESTAMPTZ DEFAULT NOW(),
        UNIQUE(user_id)
    )""",
    """CREATE TABLE IF NOT EXISTS generated_documents (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        case_id UUID REFERENCES cases(id) ON DELETE SET NULL,
        document_type TEXT NOT NULL,
        document_category TEXT NOT NULL DEFAULT 'civil',
        title TEXT,
        generated_text TEXT NOT NULL DEFAULT '',
        preview_text TEXT,
        ai_model TEXT,
        used_ai BOOLEAN NOT NULL DEFAULT true,
        has_docx_export BOOLEAN NOT NULL DEFAULT false,
        has_pdf_export BOOLEAN NOT NULL DEFAULT false,
        last_exported_at TIMESTAMPTZ,
        e_court_ready BOOLEAN NOT NULL DEFAULT false,
        filing_blockers JSONB NOT NULL DEFAULT '[]',
        created_at TIMESTAMPTZ DEFAULT NOW(),
        updated_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_generated_documents_user_id ON generated_documents(user_id)",
    "CREATE INDEX IF NOT EXISTS idx_generated_documents_created_at ON generated_documents(created_at DESC)",
    "CREATE INDEX IF NOT EXISTS idx_cases_user_id ON cases(user_id)",
    # ── AI Analyze tables ────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS contract_analyses (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        file_name TEXT,
        file_url TEXT,
        file_size BIGINT,
        contract_type TEXT,
        risk_level TEXT,
        critical_risks JSONB NOT NULL DEFAULT '[]',
        medium_risks JSONB NOT NULL DEFAULT '[]',
        ok_points JSONB NOT NULL DEFAULT '[]',
        recommendations JSONB NOT NULL DEFAULT '[]',
        summary TEXT,
        ai_model TEXT,
        tokens_used INT,
        processing_time_ms INT,
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_contract_analyses_user_id ON contract_analyses(user_id)",
    """CREATE TABLE IF NOT EXISTS document_intakes (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        case_id UUID REFERENCES cases(id) ON DELETE SET NULL,
        source_file_name TEXT,
        classified_type TEXT NOT NULL DEFAULT 'unknown',
        document_language TEXT,
        jurisdiction TEXT NOT NULL DEFAULT 'UA',
        primary_party_role TEXT,
        identified_parties JSONB NOT NULL DEFAULT '[]',
        subject_matter TEXT,
        financial_exposure_amount NUMERIC,
        financial_exposure_currency TEXT,
        financial_exposure_type TEXT,
        document_date TEXT,
        deadline_from_document TEXT,
        urgency_level TEXT,
        risk_level_legal TEXT,
        risk_level_procedural TEXT,
        risk_level_financial TEXT,
        detected_issues JSONB NOT NULL DEFAULT '[]',
        classifier_confidence NUMERIC,
        classifier_model TEXT,
        raw_text_preview TEXT,
        tags JSONB NOT NULL DEFAULT '[]',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_document_intakes_user_id ON document_intakes(user_id)",
    """CREATE TABLE IF NOT EXISTS intake_comments (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        intake_id UUID NOT NULL REFERENCES document_intakes(id) ON DELETE CASCADE,
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        content TEXT NOT NULL,
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_intake_comments_intake_id ON intake_comments(intake_id)",
    """CREATE TABLE IF NOT EXISTS knowledge_entries (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        title TEXT NOT NULL,
        content TEXT NOT NULL DEFAULT '',
        category TEXT,
        tags JSONB NOT NULL DEFAULT '[]',
        created_at TIMESTAMPTZ DEFAULT NOW(),
        updated_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_knowledge_entries_user_id ON knowledge_entries(user_id)",
    # ── Deadlines ────────────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS deadlines (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        case_id UUID REFERENCES cases(id) ON DELETE SET NULL,
        document_id UUID REFERENCES generated_documents(id) ON DELETE SET NULL,
        title TEXT NOT NULL,
        deadline_type TEXT,
        start_date TIMESTAMPTZ,
        end_date TIMESTAMPTZ,
        reminder_sent BOOLEAN NOT NULL DEFAULT false,
        notes TEXT,
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_deadlines_user_id ON deadlines(user_id)",
    """CREATE TABLE IF NOT EXISTS deadline_notifications (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        deadline_id UUID NOT NULL REFERENCES deadlines(id) ON DELETE CASCADE,
        channel TEXT NOT NULL DEFAULT 'in_app',
        message TEXT NOT NULL,
        scheduled_for TIMESTAMPTZ,
        sent_at TIMESTAMPTZ,
        status TEXT NOT NULL DEFAULT 'pending',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_deadline_notifications_user_id ON deadline_notifications(user_id)",
    # ── Case Law ─────────────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS case_law_items (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        source TEXT NOT NULL DEFAULT 'manual',
        decision_id TEXT,
        case_number TEXT,
        court_name TEXT,
        judge_name TEXT,
        decision_date DATE,
        doc_type TEXT,
        summary TEXT,
        full_text TEXT,
        relevance_score NUMERIC,
        tags JSONB NOT NULL DEFAULT '[]',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_case_law_user_id ON case_law_items(user_id)",
    "CREATE INDEX IF NOT EXISTS idx_case_law_decision_date ON case_law_items(decision_date DESC)",
    """CREATE TABLE IF NOT EXISTS case_law_digest (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        title TEXT NOT NULL DEFAULT 'Дайджест',
        summary TEXT,
        source TEXT DEFAULT 'ai',
        query TEXT,
        items_count INT NOT NULL DEFAULT 0,
        tags JSONB NOT NULL DEFAULT '[]',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_case_law_digest_user_id ON case_law_digest(user_id)",
    # ── Calculations ─────────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS calculations (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        calculation_type TEXT NOT NULL DEFAULT 'full_claim',
        title TEXT,
        input_payload JSONB NOT NULL DEFAULT '{}',
        output_payload JSONB NOT NULL DEFAULT '{}',
        notes TEXT,
        created_at TIMESTAMPTZ DEFAULT NOW(),
        updated_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_calculations_user_id ON calculations(user_id)",
    # ── Audit Logs ───────────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS audit_logs (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        action TEXT NOT NULL,
        entity_type TEXT,
        entity_id TEXT,
        metadata JSONB NOT NULL DEFAULT '{}',
        integrity_scope TEXT,
        integrity_prev_hash TEXT,
        integrity_hash TEXT,
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_audit_logs_user_id ON audit_logs(user_id)",
    "CREATE INDEX IF NOT EXISTS idx_audit_logs_created_at ON audit_logs(created_at DESC)",
    # ── Registry Watch ───────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS registry_watch_items (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        source TEXT NOT NULL DEFAULT 'opendatabot',
        registry_type TEXT NOT NULL DEFAULT 'company',
        identifier TEXT NOT NULL,
        entity_name TEXT NOT NULL DEFAULT '',
        status TEXT NOT NULL DEFAULT 'active',
        check_interval_hours INT NOT NULL DEFAULT 24,
        last_checked_at TIMESTAMPTZ,
        next_check_at TIMESTAMPTZ,
        last_change_at TIMESTAMPTZ,
        latest_snapshot JSONB,
        notes TEXT,
        created_at TIMESTAMPTZ DEFAULT NOW(),
        updated_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_registry_watch_user_id ON registry_watch_items(user_id)",
    """CREATE TABLE IF NOT EXISTS registry_events (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        watch_item_id UUID REFERENCES registry_watch_items(id) ON DELETE CASCADE,
        event_type TEXT NOT NULL,
        severity TEXT NOT NULL DEFAULT 'info',
        title TEXT NOT NULL,
        details JSONB NOT NULL DEFAULT '{}',
        observed_at TIMESTAMPTZ DEFAULT NOW(),
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_registry_events_user_id ON registry_events(user_id)",
    """CREATE TABLE IF NOT EXISTS registry_snapshots (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        watch_item_id UUID NOT NULL REFERENCES registry_watch_items(id) ON DELETE CASCADE,
        snapshot JSONB NOT NULL DEFAULT '{}',
        source TEXT NOT NULL DEFAULT 'opendatabot',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_registry_snapshots_user_id ON registry_snapshots(user_id)",
    """CREATE TABLE IF NOT EXISTS legal_sources (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        code TEXT NOT NULL,
        title TEXT NOT NULL,
        article TEXT,
        source_url TEXT,
        metadata JSONB NOT NULL DEFAULT '{}',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_legal_sources_user_id ON legal_sources(user_id)",
    """CREATE TABLE IF NOT EXISTS legal_source_chunks (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        source_id UUID NOT NULL REFERENCES legal_sources(id) ON DELETE CASCADE,
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        chunk_index INT NOT NULL DEFAULT 0,
        content TEXT NOT NULL,
        embedding JSONB NOT NULL DEFAULT '[]',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_legal_source_chunks_source_id ON legal_source_chunks(source_id)",
    """CREATE TABLE IF NOT EXISTS async_jobs (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        task_id TEXT NOT NULL UNIQUE,
        job_type TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'queued',
        payload JSONB NOT NULL DEFAULT '{}',
        result JSONB,
        error_message TEXT,
        retries INT NOT NULL DEFAULT 0,
        created_at TIMESTAMPTZ DEFAULT NOW(),
        updated_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_async_jobs_user_id ON async_jobs(user_id)",
    "CREATE INDEX IF NOT EXISTS idx_async_jobs_status ON async_jobs(status)",
    """CREATE TABLE IF NOT EXISTS dead_letter_jobs (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        task_id TEXT NOT NULL,
        job_type TEXT NOT NULL,
        error_message TEXT,
        meta JSONB NOT NULL DEFAULT '{}',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_dead_letter_jobs_user_id ON dead_letter_jobs(user_id)",
    # ── Forum ─────────────────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS forum_posts (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        case_id UUID REFERENCES cases(id) ON DELETE SET NULL,
        title TEXT NOT NULL,
        content TEXT NOT NULL DEFAULT '',
        category TEXT,
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_forum_posts_user_id ON forum_posts(user_id)",
    """CREATE TABLE IF NOT EXISTS forum_comments (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        post_id UUID NOT NULL REFERENCES forum_posts(id) ON DELETE CASCADE,
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        content TEXT NOT NULL,
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_forum_comments_post_id ON forum_comments(post_id)",
    # ── E-Court Submissions ───────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS ecourt_submissions (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        document_id UUID REFERENCES generated_documents(id) ON DELETE SET NULL,
        provider TEXT NOT NULL DEFAULT 'manual',
        external_submission_id TEXT NOT NULL,
        status TEXT NOT NULL DEFAULT 'submitted',
        court_name TEXT NOT NULL DEFAULT '',
        signer_method TEXT,
        note TEXT,
        tracking_url TEXT,
        error_message TEXT,
        submitted_at TIMESTAMPTZ DEFAULT NOW(),
        updated_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_ecourt_submissions_user_id ON ecourt_submissions(user_id)",
    # ── KEP auth ──────────────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS auth_challenges (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        nonce TEXT NOT NULL,
        provider TEXT NOT NULL DEFAULT 'local_key',
        purpose TEXT NOT NULL DEFAULT 'login',
        expires_at TIMESTAMPTZ NOT NULL,
        used_at TIMESTAMPTZ,
        ip TEXT,
        user_agent TEXT,
        origin TEXT,
        ua_hash TEXT,
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_auth_challenges_expires_at ON auth_challenges(expires_at)",
    "CREATE INDEX IF NOT EXISTS idx_auth_challenges_used_at ON auth_challenges(used_at)",
    "ALTER TABLE auth_challenges ADD COLUMN IF NOT EXISTS provider TEXT NOT NULL DEFAULT 'local_key'",
    "ALTER TABLE auth_challenges ADD COLUMN IF NOT EXISTS origin TEXT",
    "ALTER TABLE auth_challenges ADD COLUMN IF NOT EXISTS ua_hash TEXT",
    """CREATE TABLE IF NOT EXISTS user_kep_identities (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        user_id UUID NOT NULL REFERENCES users(id) ON DELETE CASCADE,
        cert_fingerprint TEXT NOT NULL UNIQUE,
        subject_dn TEXT NOT NULL DEFAULT '',
        serial_number TEXT NOT NULL DEFAULT '',
        issuer_dn TEXT NOT NULL DEFAULT '',
        valid_from TIMESTAMPTZ,
        valid_to TIMESTAMPTZ,
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_user_kep_identities_user_id ON user_kep_identities(user_id)",
    # ── Document Versions ────────────────────────────────────────────────────
    """CREATE TABLE IF NOT EXISTS document_versions (
        id UUID PRIMARY KEY DEFAULT gen_random_uuid(),
        document_id UUID NOT NULL REFERENCES generated_documents(id) ON DELETE CASCADE,
        version_number INT NOT NULL DEFAULT 1,
        action TEXT NOT NULL DEFAULT 'edit',
        generated_text TEXT NOT NULL DEFAULT '',
        created_at TIMESTAMPTZ DEFAULT NOW()
    )""",
    "CREATE INDEX IF NOT EXISTS idx_document_versions_doc_id ON document_versions(document_id)",
]

# ── Auth helpers ──────────────────────────────────────────────────────────────
_SECRET = os.getenv("SECRET_KEY", "dev-secret-change-me")
ALLOW_DEV_AUTH = os.getenv("ALLOW_DEV_AUTH", "false").lower() == "true"
_bearer = HTTPBearer(auto_error=False)
_RATE_LIMIT_BUCKETS: dict[str, deque[float]] = defaultdict(deque)
KEP_ALLOWED_PURPOSES = {"login", "link"}
KEP_ALLOWED_PROVIDERS = {"local_key", "cloud_sign"}
KEP_CHALLENGE_TTL_SECONDS = 180
KEP_ISSUED_AT_MAX_SKEW_SECONDS = 300


def _check_rate_limit(key: str, limit: int, window_seconds: int) -> None:
    now = time.time()
    bucket = _RATE_LIMIT_BUCKETS[key]
    while bucket and (now - bucket[0]) > window_seconds:
        bucket.popleft()
    if len(bucket) >= limit:
        raise HTTPException(status_code=429, detail="Забагато спроб. Спробуйте пізніше.")
    bucket.append(now)


def _make_token(user_id: str, email: str) -> str:
    """3-part token: header.payload.sig  (JWT-compatible structure for frontend decoder)."""
    header = base64.urlsafe_b64encode(b'{"alg":"HS256"}').decode().rstrip("=")
    payload_json = json.dumps({"sub": user_id, "email": email})
    payload = base64.urlsafe_b64encode(payload_json.encode()).decode().rstrip("=")
    signing_input = f"{header}.{payload}".encode()
    sig = hmac.new(_SECRET.encode(), signing_input, hashlib.sha256).hexdigest()
    return f"{header}.{payload}.{sig}"


def _decode_token(token: str) -> dict[str, str] | None:
    try:
        parts = token.split(".")
        if len(parts) == 3:
            header, payload_b64, sig = parts
            signing_input = f"{header}.{payload_b64}".encode()
            expected = hmac.new(_SECRET.encode(), signing_input, hashlib.sha256).hexdigest()
            if not hmac.compare_digest(sig, expected):
                return None
            payload_bytes = base64.urlsafe_b64decode(payload_b64 + "==")
            return json.loads(payload_bytes)
        if len(parts) == 2:
            # Legacy 2-part format (backward compat)
            payload_b64, sig = parts
            payload_bytes = base64.urlsafe_b64decode(payload_b64.encode() + b"==")
            expected = hmac.new(_SECRET.encode(), payload_bytes, hashlib.sha256).hexdigest()
            if not hmac.compare_digest(sig, expected):
                return None
            return json.loads(payload_bytes)
    except Exception:
        return None
    return None


async def _get_or_create_user(
    session: AsyncSession, email: str, name: str | None = None, password: str | None = None
) -> dict:
    row = (
        await session.execute(
            text("SELECT id, email, full_name, company, role FROM users WHERE email = :e LIMIT 1"),
            {"e": email},
        )
    ).mappings().first()
    if row:
        return dict(row)
    uid = str(uuid.uuid4())
    pw_hash = hashlib.sha256(password.encode()).hexdigest() if password else ""
    await session.execute(
        text("""
            INSERT INTO users (id, email, full_name, role, password_hash)
            VALUES (:id, :email, :name, 'user', :pw)
            ON CONFLICT (email) DO NOTHING
        """),
        {"id": uid, "email": email, "name": name or email.split("@")[0], "pw": pw_hash},
    )
    await session.commit()
    row = (
        await session.execute(
            text("SELECT id, email, full_name, company, role FROM users WHERE email = :e LIMIT 1"),
            {"e": email},
        )
    ).mappings().first()
    return dict(row) if row else {"id": uid, "email": email, "full_name": name, "role": "user"}


async def get_current_user(
    creds: Optional[HTTPAuthorizationCredentials] = Depends(_bearer),
    session: AsyncSession = Depends(get_session),
) -> dict:
    token_str = creds.credentials if creds else ""

    # Dev-mode: accept legacy offline tokens (dev-token, dev-token-<slug>)
    if ALLOW_DEV_AUTH and (not token_str or token_str.startswith("dev-token")):
        if token_str.startswith("dev-token-"):
            slug = token_str[len("dev-token-"):]          # e.g. "dev-john"
            # strip leading "dev-" prefix that frontend adds
            slug = slug[4:] if slug.startswith("dev-") else slug
            email = f"{slug}@dev.local"
            name = slug.split("-")[0].capitalize()
        else:
            email = "dev@legal-ai.local"
            name = "Dev User"
        return await _get_or_create_user(session, email, name)

    if creds is None:
        raise HTTPException(status_code=401, detail="Missing token")

    data = _decode_token(creds.credentials)
    if not data:
        raise HTTPException(status_code=401, detail="Invalid token")

    row = (
        await session.execute(
            text("SELECT id, email, full_name, company, role FROM users WHERE id = :id LIMIT 1"),
            {"id": data["sub"]},
        )
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=401, detail="User not found")
    return dict(row)


# ── Password helpers ──────────────────────────────────────────────────────────
try:
    from passlib.context import CryptContext
    _pwd_ctx = CryptContext(schemes=["bcrypt"], deprecated="auto")
    def _hash_pw(pw: str) -> str: return _pwd_ctx.hash(pw)
    def _verify_pw(pw: str, h: str) -> bool: return _pwd_ctx.verify(pw, h)
except Exception:
    def _hash_pw(pw: str) -> str: return hashlib.sha256(pw.encode()).hexdigest()
    def _verify_pw(pw: str, h: str) -> bool: return _hash_pw(pw) == h


# ============================================================================
# HEALTH
# ============================================================================
@app.get("/health")
async def health():
    return {"status": "ok", "version": "1.0.0"}


# ============================================================================
# AUTH
# ============================================================================
class AuthRequest(BaseModel):
    email: str
    password: str
    full_name: str | None = None


class KepChallengeRequest(BaseModel):
    provider: str = "local_key"
    purpose: str = "login"


class KepVerifyRequest(BaseModel):
    challenge_id: str
    signature: str
    signed_payload: str
    certificate: str
    provider: str = "local_key"


def _client_ip(raw: str | None) -> str:
    if not raw:
        return "unknown"
    # X-Forwarded-For may contain a chain: "ip1, ip2"
    return raw.split(",")[0].strip()


def _make_nonce() -> str:
    return base64.urlsafe_b64encode(os.urandom(32)).decode().rstrip("=")


def _safe_b64_decode(value: str) -> bytes:
    padded = value + ("=" * (-len(value) % 4))
    try:
        return base64.urlsafe_b64decode(padded.encode())
    except Exception:
        return base64.b64decode(padded.encode())


def _ua_hash(user_agent: str) -> str:
    return hashlib.sha256((user_agent or "").encode("utf-8")).hexdigest()


def _origin_from_request(request: Request) -> str:
    origin = request.headers.get("origin", "").strip()
    if origin:
        return origin
    host = request.headers.get("host", "").strip()
    if host:
        proto = request.headers.get("x-forwarded-proto", "https").strip() or "https"
        return f"{proto}://{host}"
    return "unknown"


def _decode_signed_payload(payload: str) -> dict:
    try:
        decoded = _safe_b64_decode(payload).decode("utf-8")
        data = json.loads(decoded)
        if isinstance(data, dict):
            return data
    except Exception:
        pass
    return {}


async def _verify_kep_with_provider(
    body: KepVerifyRequest,
    nonce: str,
    challenge_id: str,
    expected_origin: str,
    expected_ua_hash: str,
    expected_purpose: str,
) -> tuple[bool, dict]:
    payload = _decode_signed_payload(body.signed_payload)
    if payload.get("nonce") != nonce or payload.get("challenge_id") != challenge_id:
        return False, {"reason": "payload_nonce_mismatch"}
    if payload.get("origin") != expected_origin:
        return False, {"reason": "payload_origin_mismatch"}
    if payload.get("ua_hash") != expected_ua_hash:
        return False, {"reason": "payload_ua_mismatch"}
    if payload.get("purpose") != expected_purpose:
        return False, {"reason": "payload_purpose_mismatch"}
    issued_at_raw = str(payload.get("issued_at") or "").strip()
    try:
        issued_at = datetime.fromisoformat(issued_at_raw.replace("Z", "+00:00"))
        if issued_at.tzinfo is not None:
            issued_at = issued_at.astimezone(timezone.utc).replace(tzinfo=None)
    except Exception:
        return False, {"reason": "payload_issued_at_invalid"}
    if abs((datetime.utcnow() - issued_at).total_seconds()) > KEP_ISSUED_AT_MAX_SKEW_SECONDS:
        return False, {"reason": "payload_issued_at_skew"}
    if not body.signature or not body.certificate:
        return False, {"reason": "empty_signature_or_certificate"}

    # Optional external verifier (recommended for production)
    verifier_url = os.getenv("KEP_VERIFY_URL", "").strip()
    verifier_token = os.getenv("KEP_VERIFY_TOKEN", "").strip()
    if verifier_url:
        try:
            headers = {"Content-Type": "application/json"}
            if verifier_token:
                headers["Authorization"] = f"Bearer {verifier_token}"
            async with httpx.AsyncClient(timeout=20) as client:
                response = await client.post(
                    verifier_url,
                    headers=headers,
                    json={
                        "signature": body.signature,
                        "signed_payload": body.signed_payload,
                        "certificate": body.certificate,
                        "provider": body.provider,
                    },
                )
            if response.is_success:
                data = response.json() if response.content else {}
                return bool(data.get("valid", False)), data if isinstance(data, dict) else {}
            return False, {"reason": f"provider_http_{response.status_code}"}
        except Exception as e:
            return False, {"reason": f"provider_error:{e}"}

    # Local fallback/dev verification (not a cryptographic KEP validation).
    if ALLOW_DEV_AUTH and body.signature and body.certificate:
        return True, {"valid": True, "subject_dn": "CN=Dev KEP User", "serial_number": "dev-serial", "issuer_dn": "CN=Dev CA"}
    return False, {"reason": "no_verifier_configured"}


@app.post("/api/auth/login")
async def login(body: AuthRequest, session: AsyncSession = Depends(get_session)):
    row = (
        await session.execute(
            text("SELECT id, email, full_name, password_hash FROM users WHERE email = :e LIMIT 1"),
            {"e": body.email},
        )
    ).mappings().first()

    if row is None:
        if ALLOW_DEV_AUTH:
            user = await _get_or_create_user(session, body.email, body.email.split("@")[0], body.password)
            token = _make_token(str(user["id"]), body.email)
            return {"access_token": token, "token_type": "bearer",
                    "user": {"id": str(user["id"]), "email": body.email, "name": user.get("full_name")}}
        raise HTTPException(status_code=401, detail="Неправильний email або пароль")

    pw_hash = row.get("password_hash") or ""
    if pw_hash and not _verify_pw(body.password, pw_hash) and not ALLOW_DEV_AUTH:
        raise HTTPException(status_code=401, detail="Неправильний email або пароль")

    token = _make_token(str(row["id"]), body.email)
    return {
        "access_token": token, "token_type": "bearer",
        "user": {"id": str(row["id"]), "email": body.email, "name": row.get("full_name")},
    }


@app.post("/api/auth/register")
async def register(body: AuthRequest, session: AsyncSession = Depends(get_session)):
    exists = (
        await session.execute(text("SELECT 1 FROM users WHERE email = :e LIMIT 1"), {"e": body.email})
    ).first()
    if exists:
        raise HTTPException(status_code=409, detail="Користувач з таким email вже існує")

    uid = str(uuid.uuid4())
    pw_hash = _hash_pw(body.password) if body.password else ""
    await session.execute(
        text("""
            INSERT INTO users (id, email, full_name, role, password_hash)
            VALUES (:id, :email, :name, 'user', :pw)
        """),
        {"id": uid, "email": body.email, "name": body.full_name or body.email.split("@")[0], "pw": pw_hash},
    )
    await session.commit()
    token = _make_token(uid, body.email)
    return {
        "access_token": token, "token_type": "bearer",
        "user": {"id": uid, "email": body.email, "name": body.full_name},
    }


@app.post("/api/auth/kep/challenge")
async def create_kep_challenge(
    body: KepChallengeRequest,
    request: Request,
    session: AsyncSession = Depends(get_session),
):
    if body.purpose not in KEP_ALLOWED_PURPOSES:
        raise HTTPException(status_code=400, detail={"error_code": "KEP_PURPOSE_INVALID", "message": "Непідтримуваний purpose."})
    if body.provider not in KEP_ALLOWED_PROVIDERS:
        raise HTTPException(status_code=400, detail={"error_code": "KEP_PROVIDER_INVALID", "message": "Непідтримуваний провайдер КЕП."})
    ip = _client_ip(request.headers.get("x-forwarded-for") or request.client.host if request.client else "")
    _check_rate_limit(f"kep:challenge:ip:{ip}", limit=10, window_seconds=60)
    nonce = _make_nonce()
    challenge_id = str(uuid.uuid4())
    expires_at = datetime.utcnow() + timedelta(seconds=KEP_CHALLENGE_TTL_SECONDS)
    ua = request.headers.get("user-agent", "")
    origin = _origin_from_request(request)
    ua_digest = _ua_hash(ua)
    issued_at = datetime.utcnow().replace(microsecond=0).isoformat() + "Z"
    challenge_payload = {
        "challenge_id": challenge_id,
        "nonce": nonce,
        "purpose": body.purpose,
        "origin": origin,
        "ua_hash": ua_digest,
        "issued_at": issued_at,
    }
    await session.execute(
        text("""
            INSERT INTO auth_challenges (id, nonce, provider, purpose, expires_at, ip, user_agent, origin, ua_hash)
            VALUES (:id, :nonce, :provider, :purpose, :expires_at, :ip, :ua, :origin, :ua_hash)
        """),
        {
            "id": challenge_id,
            "nonce": nonce,
            "provider": body.provider,
            "purpose": body.purpose,
            "expires_at": expires_at,
            "ip": ip,
            "ua": ua,
            "origin": origin,
            "ua_hash": ua_digest,
        },
    )
    await _audit_log(session, "system", "kep_challenge_created", "auth_challenge", challenge_id, {"provider": body.provider, "purpose": body.purpose, "ip": ip})
    await session.commit()
    return {
        "challenge_id": challenge_id,
        "nonce": nonce,
        "expires_at": expires_at.isoformat() + "Z",
        "algorithms": ["RSA-PSS-SHA256", "ECDSA-SHA256"],
        "challenge_payload": challenge_payload,
    }


@app.post("/api/auth/kep/verify")
async def verify_kep_auth(
    body: KepVerifyRequest,
    request: Request,
    session: AsyncSession = Depends(get_session),
):
    ip = _client_ip(request.headers.get("x-forwarded-for") or request.client.host if request.client else "")
    _check_rate_limit(f"kep:verify:ip:{ip}", limit=5, window_seconds=60)

    cert_fingerprint = hashlib.sha256(body.certificate.encode()).hexdigest()
    _check_rate_limit(f"kep:verify:fpr:{cert_fingerprint}", limit=5, window_seconds=600)

    row = (
        await session.execute(
            text("""
                SELECT id, nonce, provider, purpose, expires_at, used_at, origin, ua_hash
                FROM auth_challenges
                WHERE id = :id
                LIMIT 1
            """),
            {"id": body.challenge_id},
        )
    ).mappings().first()
    if not row:
        await _audit_log(session, "system", "kep_verify_failed", "auth_challenge", body.challenge_id, {"error": "challenge_not_found", "ip": ip})
        await session.commit()
        raise HTTPException(status_code=401, detail={"error_code": "KEP_CHALLENGE_NOT_FOUND", "message": "Challenge не знайдено."})
    if row.get("used_at"):
        await _audit_log(session, "system", "kep_verify_failed", "auth_challenge", body.challenge_id, {"error": "challenge_replay", "ip": ip})
        await session.commit()
        raise HTTPException(status_code=401, detail={"error_code": "KEP_REPLAY_DETECTED", "message": "Challenge вже був використаний."})
    if row.get("expires_at") and row["expires_at"] < datetime.utcnow():
        await _audit_log(session, "system", "kep_verify_failed", "auth_challenge", body.challenge_id, {"error": "challenge_expired", "ip": ip})
        await session.commit()
        raise HTTPException(status_code=401, detail={"error_code": "KEP_CHALLENGE_EXPIRED", "message": "Challenge протермінований."})
    if str(row.get("provider") or "local_key") != body.provider:
        await _audit_log(session, "system", "kep_verify_failed", "auth_challenge", body.challenge_id, {"error": "provider_mismatch", "ip": ip})
        await session.commit()
        raise HTTPException(status_code=401, detail={"error_code": "KEP_PROVIDER_MISMATCH", "message": "Challenge створено для іншого КЕП-провайдера."})
    if str(row.get("purpose") or "login") != "login":
        await _audit_log(session, "system", "kep_verify_failed", "auth_challenge", body.challenge_id, {"error": "purpose_mismatch", "ip": ip})
        await session.commit()
        raise HTTPException(status_code=401, detail={"error_code": "KEP_PURPOSE_MISMATCH", "message": "Невірний тип challenge для входу."})

    ok, provider_data = await _verify_kep_with_provider(
        body,
        str(row["nonce"]),
        body.challenge_id,
        str(row.get("origin") or _origin_from_request(request)),
        str(row.get("ua_hash") or _ua_hash(request.headers.get("user-agent", ""))),
        "login",
    )
    if not ok:
        await _audit_log(session, "system", "kep_verify_failed", "auth_challenge", body.challenge_id, {"error": provider_data.get("reason", "verify_failed"), "ip": ip})
        await session.commit()
        raise HTTPException(status_code=401, detail={"error_code": "KEP_VERIFY_FAILED", "message": "Не вдалося перевірити КЕП підпис."})

    subject_dn = str(provider_data.get("subject_dn") or "CN=KEP User")
    serial_number = str(provider_data.get("serial_number") or cert_fingerprint[:16])
    issuer_dn = str(provider_data.get("issuer_dn") or "CN=Unknown CA")
    cert_email = f"kep-{cert_fingerprint[:12]}@local"

    identity = (
        await session.execute(
            text("""
                SELECT user_id FROM user_kep_identities
                WHERE cert_fingerprint = :fp
                LIMIT 1
            """),
            {"fp": cert_fingerprint},
        )
    ).mappings().first()

    if identity:
        user_row = (
            await session.execute(
                text("SELECT id, email, full_name FROM users WHERE id = :uid LIMIT 1"),
                {"uid": identity["user_id"]},
            )
        ).mappings().first()
    else:
        user = await _get_or_create_user(session, cert_email, subject_dn.replace("CN=", "")[:120], password=None)
        user_row = {"id": user["id"], "email": user.get("email", cert_email), "full_name": user.get("full_name")}
        await session.execute(
            text("""
                INSERT INTO user_kep_identities
                    (user_id, cert_fingerprint, subject_dn, serial_number, issuer_dn, valid_from, valid_to)
                VALUES
                    (:uid, :fp, :subject_dn, :serial_number, :issuer_dn, NOW(), NOW() + INTERVAL '365 days')
                ON CONFLICT (cert_fingerprint) DO NOTHING
            """),
            {
                "uid": str(user_row["id"]),
                "fp": cert_fingerprint,
                "subject_dn": subject_dn,
                "serial_number": serial_number,
                "issuer_dn": issuer_dn,
            },
        )

    await session.execute(
        text("UPDATE auth_challenges SET used_at = NOW() WHERE id = :id"),
        {"id": body.challenge_id},
    )
    token = _make_token(str(user_row["id"]), str(user_row.get("email") or cert_email))
    await _audit_log(session, str(user_row["id"]), "kep_verify_ok", "auth_challenge", body.challenge_id, {"provider": body.provider, "cert_fingerprint": cert_fingerprint, "ip": ip})
    await session.commit()
    return {
        "access_token": token,
        "token_type": "bearer",
        "user": {
            "id": str(user_row["id"]),
            "email": str(user_row.get("email") or cert_email),
            "name": str(user_row.get("full_name") or "KEP User"),
        },
    }


@app.post("/api/auth/kep/link")
async def link_kep_identity(
    body: KepVerifyRequest,
    request: Request,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    ip = _client_ip(request.headers.get("x-forwarded-for") or request.client.host if request.client else "")
    _check_rate_limit(f"kep:link:ip:{ip}", limit=5, window_seconds=60)
    cert_fingerprint = hashlib.sha256(body.certificate.encode()).hexdigest()
    _check_rate_limit(f"kep:link:fpr:{cert_fingerprint}", limit=5, window_seconds=600)
    row = (
        await session.execute(
            text("SELECT id, nonce, provider, purpose, expires_at, used_at, origin, ua_hash FROM auth_challenges WHERE id = :id LIMIT 1"),
            {"id": body.challenge_id},
        )
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=401, detail={"error_code": "KEP_CHALLENGE_NOT_FOUND", "message": "Challenge не знайдено."})
    if row.get("used_at") or (row.get("expires_at") and row["expires_at"] < datetime.utcnow()):
        raise HTTPException(status_code=401, detail={"error_code": "KEP_CHALLENGE_INVALID", "message": "Challenge недійсний."})
    if str(row.get("provider") or "local_key") != body.provider:
        raise HTTPException(status_code=401, detail={"error_code": "KEP_PROVIDER_MISMATCH", "message": "Challenge створено для іншого КЕП-провайдера."})
    if str(row.get("purpose") or "login") != "link":
        raise HTTPException(status_code=401, detail={"error_code": "KEP_PURPOSE_MISMATCH", "message": "Невірний тип challenge для прив'язки."})
    ok, provider_data = await _verify_kep_with_provider(
        body,
        str(row["nonce"]),
        body.challenge_id,
        str(row.get("origin") or _origin_from_request(request)),
        str(row.get("ua_hash") or _ua_hash(request.headers.get("user-agent", ""))),
        "link",
    )
    if not ok:
        await _audit_log(session, str(current_user["id"]), "kep_verify_failed", "auth_challenge", body.challenge_id, {"error": provider_data.get("reason", "verify_failed"), "ip": ip})
        await session.commit()
        raise HTTPException(status_code=401, detail={"error_code": "KEP_VERIFY_FAILED", "message": "Не вдалося перевірити КЕП підпис."})

    await session.execute(
        text("""
            INSERT INTO user_kep_identities
                (user_id, cert_fingerprint, subject_dn, serial_number, issuer_dn, valid_from, valid_to)
            VALUES
                (:uid, :fp, :subject_dn, :serial_number, :issuer_dn, NOW(), NOW() + INTERVAL '365 days')
            ON CONFLICT (cert_fingerprint)
            DO UPDATE SET user_id = EXCLUDED.user_id
        """),
        {
            "uid": str(current_user["id"]),
            "fp": cert_fingerprint,
            "subject_dn": str(provider_data.get("subject_dn") or "CN=KEP User"),
            "serial_number": str(provider_data.get("serial_number") or cert_fingerprint[:16]),
            "issuer_dn": str(provider_data.get("issuer_dn") or "CN=Unknown CA"),
        },
    )
    await session.execute(text("UPDATE auth_challenges SET used_at = NOW() WHERE id = :id"), {"id": body.challenge_id})
    await _audit_log(session, str(current_user["id"]), "kep_identity_linked", "auth_challenge", body.challenge_id, {"cert_fingerprint": cert_fingerprint, "ip": ip})
    await session.commit()
    return {"status": "linked", "cert_fingerprint": cert_fingerprint}


@app.get("/api/auth/me")
async def get_me(current_user: dict = Depends(get_current_user)):
    return {
        "user_id": str(current_user["id"]),
        "email": current_user["email"],
        "full_name": current_user.get("full_name"),
        "company": current_user.get("company"),
        "role": current_user.get("role", "user"),
        "workspace_id": str(current_user["id"]),
    }


class UpdateMeRequest(BaseModel):
    full_name: str | None = None
    company: str | None = None
    entity_type: str | None = None
    tax_id: str | None = None
    address: str | None = None
    phone: str | None = None
    logo_url: str | None = None


@app.patch("/api/auth/me")
async def update_me(
    body: UpdateMeRequest,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    fields = {k: v for k, v in body.model_dump().items() if v is not None}
    if fields:
        sets = ", ".join(f"{k} = :{k}" for k in fields)
        fields["uid"] = str(current_user["id"])
        await session.execute(text(f"UPDATE users SET {sets} WHERE id = :uid"), fields)
        await session.commit()
    return {"status": "ok"}


# ============================================================================
# BILLING
# ============================================================================
_PLANS = [
    {"id": "FREE",     "name": "FREE",     "price_uah": 0,    "docs_limit": 5,   "features": ["Генерація: 5/міс"]},
    {"id": "PRO",      "name": "PRO",      "price_uah": 499,  "docs_limit": 50,  "features": ["Генерація: 50/міс", "Судова практика"]},
    {"id": "PRO_PLUS", "name": "PRO+",     "price_uah": 999,  "docs_limit": None,"features": ["Безлімітна генерація", "Моніторинг", "Е-Суд"]},
]


@app.get("/api/billing/plans")
async def get_plans():
    return {"items": _PLANS}


@app.get("/api/billing/subscription")
async def get_subscription(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    row = (
        await session.execute(
            text("SELECT plan, docs_used, docs_limit FROM subscriptions WHERE user_id = :uid LIMIT 1"),
            {"uid": uid},
        )
    ).mappings().first()

    if row is None:
        # Create default FREE subscription
        try:
            await session.execute(
                text("INSERT INTO subscriptions (user_id, plan, docs_used, docs_limit) VALUES (:uid, 'FREE', 0, 5) ON CONFLICT (user_id) DO NOTHING"),
                {"uid": uid},
            )
            await session.commit()
        except Exception:
            pass
        plan, docs_used, docs_limit = "FREE", 0, 5
    else:
        plan = row["plan"]
        docs_used = row["docs_used"]
        docs_limit = row["docs_limit"]

    return {
        "plan": plan,
        "status": "active",
        "usage": {
            "docs_used": docs_used,
            "docs_limit": docs_limit,
        },
    }


@app.post("/api/billing/subscribe")
async def subscribe_plan(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    plan = body.get("plan", "FREE")
    limits = {"FREE": 5, "PRO": 50, "PRO_PLUS": None}
    docs_limit = limits.get(plan, 5)
    uid = str(current_user["id"])
    await session.execute(
        text("""
            INSERT INTO subscriptions (user_id, plan, docs_used, docs_limit)
            VALUES (:uid, :plan, 0, :lim)
            ON CONFLICT (user_id) DO UPDATE SET plan = :plan, docs_limit = :lim
        """),
        {"uid": uid, "plan": plan, "lim": docs_limit},
    )
    await session.commit()
    return {"status": "ok", "plan": plan}


class LegalSourceIndexRequest(BaseModel):
    code: str
    title: str
    article: str | None = None
    text: str
    source_url: str | None = None
    metadata: dict[str, Any] | None = None


class LegalRetrieveRequest(BaseModel):
    query: str
    top_k: int = 5


def _text_to_embedding(text_value: str, dim: int = 64) -> list[float]:
    # Deterministic lightweight embedding fallback (works without extra deps).
    buckets = [0.0] * dim
    for token in re.findall(r"[\\w\\-]+", (text_value or "").lower()):
        idx = abs(hash(token)) % dim
        buckets[idx] += 1.0
    norm = sum(v * v for v in buckets) ** 0.5
    if norm > 0:
        buckets = [v / norm for v in buckets]
    return buckets


def _vector_backend_info() -> dict[str, Any]:
    backend = os.getenv("LEGAL_VECTOR_BACKEND", "pgvector").lower()
    if backend not in {"pgvector", "chroma"}:
        backend = "pgvector"
    return {
        "backend": backend,
        "ready": True,
        "note": "Storage is persisted in Postgres tables; adapter target can be pgvector/chroma via LEGAL_VECTOR_BACKEND.",
    }


def _cosine(a: list[float], b: list[float]) -> float:
    if not a or not b:
        return 0.0
    size = min(len(a), len(b))
    dot = sum(a[i] * b[i] for i in range(size))
    return float(dot)


def _chunk_text(raw: str, chunk_size: int = 700, overlap: int = 120) -> list[str]:
    text_value = (raw or "").strip()
    if not text_value:
        return []
    chunks: list[str] = []
    step = max(1, chunk_size - overlap)
    for start in range(0, len(text_value), step):
        chunk = text_value[start:start + chunk_size].strip()
        if chunk:
            chunks.append(chunk)
        if start + chunk_size >= len(text_value):
            break
    return chunks


async def _retrieve_legal_chunks(session: AsyncSession, uid: str, query: str, top_k: int = 5) -> list[dict[str, Any]]:
    query_vec = _text_to_embedding(query)
    rows = (
        await session.execute(
            text(
                """
                SELECT c.id, c.content, c.embedding, c.chunk_index,
                       s.code, s.title, s.article, s.source_url
                FROM legal_source_chunks c
                JOIN legal_sources s ON s.id = c.source_id
                WHERE c.user_id = :uid
                LIMIT 2000
                """
            ),
            {"uid": uid},
        )
    ).mappings().all()
    ranked: list[tuple[float, dict[str, Any]]] = []
    for r in rows:
        embedding = r.get("embedding") or []
        if isinstance(embedding, str):
            try:
                embedding = json.loads(embedding)
            except Exception:
                embedding = []
        score = _cosine(query_vec, embedding if isinstance(embedding, list) else [])
        payload = dict(r)
        payload["score"] = round(score, 6)
        ranked.append((score, payload))
    ranked.sort(key=lambda x: x[0], reverse=True)
    return [item for _, item in ranked[: max(1, min(top_k, 20))]]


def _build_citations(chunks: list[dict[str, Any]]) -> list[dict[str, Any]]:
    citations: list[dict[str, Any]] = []
    for idx, chunk in enumerate(chunks, start=1):
        locator = chunk.get("article") or f"chunk#{chunk.get('chunk_index', 0)}"
        citations.append(
            {
                "citation_id": f"CIT-{idx}",
                "code": chunk.get("code"),
                "title": chunk.get("title"),
                "locator": locator,
                "source_url": chunk.get("source_url"),
                "evidence_span": (chunk.get("content") or "")[:240],
                "score": chunk.get("score", 0.0),
            }
        )
    return citations


def _grounding_quality_metrics(generated_text: str, citations: list[dict[str, Any]]) -> dict[str, Any]:
    text_u = (generated_text or "").upper()
    if not citations:
        return {"citation_coverage": 0.0, "faithfulness": 0.0, "citations_count": 0}
    mentioned = 0
    faithful = 0
    for c in citations:
        locator = str(c.get("locator") or "").upper()
        title = str(c.get("title") or "").upper()
        evidence = str(c.get("evidence_span") or "")[:80].upper()
        if locator and locator in text_u or (title and title in text_u):
            mentioned += 1
        if evidence and evidence[:30] in text_u:
            faithful += 1
    total = len(citations)
    return {
        "citation_coverage": round(mentioned / total, 4),
        "faithfulness": round(faithful / total, 4),
        "citations_count": total,
    }


@app.post("/api/legal-brain/sources/index")
async def index_legal_source(
    body: LegalSourceIndexRequest,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    source_id = str(uuid.uuid4())
    await session.execute(
        text(
            """
            INSERT INTO legal_sources (id, user_id, code, title, article, source_url, metadata)
            VALUES (:id, :uid, :code, :title, :article, :source_url, :metadata::jsonb)
            """
        ),
        {
            "id": source_id,
            "uid": uid,
            "code": body.code,
            "title": body.title,
            "article": body.article,
            "source_url": body.source_url,
            "metadata": json.dumps(body.metadata or {}, ensure_ascii=False),
        },
    )
    chunks = _chunk_text(body.text)
    for idx, chunk in enumerate(chunks):
        await session.execute(
            text(
                """
                INSERT INTO legal_source_chunks (id, source_id, user_id, chunk_index, content, embedding)
                VALUES (:id, :source_id, :uid, :chunk_index, :content, :embedding::jsonb)
                """
            ),
            {
                "id": str(uuid.uuid4()),
                "source_id": source_id,
                "uid": uid,
                "chunk_index": idx,
                "content": chunk,
                "embedding": json.dumps(_text_to_embedding(chunk), ensure_ascii=False),
            },
        )
    await session.commit()
    return {"status": "indexed", "source_id": source_id, "chunks": len(chunks)}


@app.post("/api/legal-brain/retrieve")
async def retrieve_legal_context(
    body: LegalRetrieveRequest,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    chunks = await _retrieve_legal_chunks(session, uid, body.query, body.top_k)
    return {"query": body.query, "items": _build_citations(chunks), "vector_backend": _vector_backend_info()}


@app.get("/api/legal-brain/vector-backend")
async def get_vector_backend_status():
    return _vector_backend_info()


# ============================================================================
# DOCUMENTS
# ============================================================================
_DOC_CATEGORIES: dict[str, str] = {
    "pozov_do_sudu": "civil", "pozov_trudovyi": "labor",
    "appeal_complaint": "civil", "dohovir_kupivli_prodazhu": "contract",
    "dohovir_orendi": "contract", "dohovir_nadannia_posluh": "contract",
    "pretenziya": "civil", "dovirennist": "civil",
}

_DOC_LABELS: dict[str, str] = {
    "pozov_do_sudu": "Позов до суду",
    "pozov_trudovyi": "Трудовий позов",
    "appeal_complaint": "Апеляційна скарга",
    "dohovir_kupivli_prodazhu": "Договір купівлі-продажу",
    "dohovir_orendi": "Договір оренди",
    "dohovir_nadannia_posluh": "Договір про надання послуг",
    "pretenziya": "Претензія",
    "dovirennist": "Довіреність",
}

_FORM_SCHEMA_CATALOG: dict[str, dict[str, Any]] = {
    "pozov_do_sudu": {
        "title": "Схема для позову до суду",
        "fields": [
            {"name": "court_name", "label": "Суд", "type": "text", "required": True},
            {"name": "claimant_name", "label": "Позивач", "type": "text", "required": True},
            {"name": "respondent_name", "label": "Відповідач", "type": "text", "required": True},
            {"name": "claim_subject", "label": "Суть позову", "type": "textarea", "required": True},
            {"name": "claim_amount", "label": "Ціна позову", "type": "number", "required": False},
            {"name": "facts", "label": "Обставини", "type": "textarea", "required": True},
            {"name": "legal_basis", "label": "Правова підстава", "type": "textarea", "required": True},
            {"name": "requests", "label": "Прохальна частина", "type": "textarea", "required": True},
        ],
    },
    "pozov_trudovyi": {
        "title": "Схема для трудового позову",
        "fields": [
            {"name": "court_name", "label": "Суд", "type": "text", "required": True},
            {"name": "employee_name", "label": "Працівник", "type": "text", "required": True},
            {"name": "employer_name", "label": "Роботодавець", "type": "text", "required": True},
            {"name": "employment_period", "label": "Період роботи", "type": "text", "required": True},
            {"name": "violation_description", "label": "Порушення", "type": "textarea", "required": True},
            {"name": "claim_amount", "label": "Сума вимог", "type": "number", "required": False},
            {"name": "requests", "label": "Прохальна частина", "type": "textarea", "required": True},
        ],
    },
    "appeal_complaint": {
        "title": "Схема для апеляційної скарги",
        "fields": [
            {"name": "appeal_court", "label": "Апеляційний суд", "type": "text", "required": True},
            {"name": "first_instance_court", "label": "Суд першої інстанції", "type": "text", "required": True},
            {"name": "case_number", "label": "Номер справи", "type": "text", "required": True},
            {"name": "appellant_name", "label": "Апелянт", "type": "text", "required": True},
            {"name": "contested_decision_date", "label": "Дата рішення", "type": "date", "required": True},
            {"name": "appeal_grounds", "label": "Підстави оскарження", "type": "textarea", "required": True},
            {"name": "requests", "label": "Вимоги апеляції", "type": "textarea", "required": True},
        ],
    },
    "dohovir_kupivli_prodazhu": {
        "title": "Схема для договору купівлі-продажу",
        "fields": [
            {"name": "seller_name", "label": "Продавець", "type": "text", "required": True},
            {"name": "buyer_name", "label": "Покупець", "type": "text", "required": True},
            {"name": "subject", "label": "Предмет договору", "type": "textarea", "required": True},
            {"name": "price", "label": "Ціна", "type": "number", "required": True},
            {"name": "payment_terms", "label": "Умови оплати", "type": "textarea", "required": True},
            {"name": "delivery_terms", "label": "Умови передачі", "type": "textarea", "required": False},
        ],
    },
    "dohovir_orendi": {
        "title": "Схема для договору оренди",
        "fields": [
            {"name": "lessor_name", "label": "Орендодавець", "type": "text", "required": True},
            {"name": "lessee_name", "label": "Орендар", "type": "text", "required": True},
            {"name": "rent_object", "label": "Об’єкт оренди", "type": "textarea", "required": True},
            {"name": "rent_amount", "label": "Орендна плата", "type": "number", "required": True},
            {"name": "term", "label": "Строк оренди", "type": "text", "required": True},
            {"name": "termination_terms", "label": "Порядок розірвання", "type": "textarea", "required": False},
        ],
    },
    "dohovir_nadannia_posluh": {
        "title": "Схема для договору послуг",
        "fields": [
            {"name": "provider_name", "label": "Виконавець", "type": "text", "required": True},
            {"name": "customer_name", "label": "Замовник", "type": "text", "required": True},
            {"name": "service_scope", "label": "Перелік послуг", "type": "textarea", "required": True},
            {"name": "price", "label": "Вартість", "type": "number", "required": True},
            {"name": "deadline", "label": "Строк виконання", "type": "text", "required": False},
            {"name": "acceptance_terms", "label": "Порядок приймання", "type": "textarea", "required": False},
        ],
    },
    "pretenziya": {
        "title": "Схема для претензії",
        "fields": [
            {"name": "recipient_name", "label": "Адресат", "type": "text", "required": True},
            {"name": "sender_name", "label": "Заявник", "type": "text", "required": True},
            {"name": "contract_reference", "label": "Посилання на договір", "type": "text", "required": False},
            {"name": "violation_description", "label": "Опис порушення", "type": "textarea", "required": True},
            {"name": "claim_amount", "label": "Сума претензії", "type": "number", "required": False},
            {"name": "deadline_to_respond", "label": "Строк для відповіді", "type": "date", "required": False},
            {"name": "requests", "label": "Вимоги", "type": "textarea", "required": True},
        ],
    },
    "dovirennist": {
        "title": "Схема для довіреності",
        "fields": [
            {"name": "principal_name", "label": "Довіритель", "type": "text", "required": True},
            {"name": "agent_name", "label": "Представник", "type": "text", "required": True},
            {"name": "powers", "label": "Повноваження", "type": "textarea", "required": True},
            {"name": "valid_until", "label": "Строк дії", "type": "date", "required": False},
            {"name": "place", "label": "Місце видачі", "type": "text", "required": False},
        ],
    },
}


@app.get("/api/documents/types")
async def get_document_types():
    return [
        {"doc_type": k, "title": v, "category": _DOC_CATEGORIES.get(k, "civil"), "procedure": "general"}
        for k, v in _DOC_LABELS.items()
    ]


@app.get("/api/documents/history")
async def get_documents_history(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    query: str | None = None,
    doc_type: str | None = None,
    case_id: str | None = None,
    sort_by: str = "created_at",
    sort_dir: str = "desc",
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    conditions = ["user_id = :uid"]
    params: dict[str, Any] = {"uid": uid}

    if query:
        conditions.append("(title ILIKE :q OR preview_text ILIKE :q)")
        params["q"] = f"%{query}%"
    if doc_type:
        conditions.append("document_type = :dt")
        params["dt"] = doc_type
    if case_id:
        conditions.append("case_id = :cid")
        params["cid"] = case_id

    where = " AND ".join(conditions)
    order = f"{'created_at' if sort_by not in ('created_at', 'document_type') else sort_by} {'DESC' if sort_dir == 'desc' else 'ASC'}"

    total_row = (await session.execute(text(f"SELECT COUNT(*) FROM generated_documents WHERE {where}"), params)).scalar()
    total = total_row or 0

    params["limit"] = page_size
    params["offset"] = (page - 1) * page_size
    rows = (
        await session.execute(
            text(f"""
                SELECT id, document_type, document_category, title, preview_text,
                       ai_model, used_ai, has_docx_export, has_pdf_export,
                       last_exported_at, e_court_ready, filing_blockers, case_id, created_at
                FROM generated_documents WHERE {where}
                ORDER BY {order} LIMIT :limit OFFSET :offset
            """),
            params,
        )
    ).mappings().all()

    sub_row = (
        await session.execute(
            text("SELECT docs_used, docs_limit FROM subscriptions WHERE user_id = :uid LIMIT 1"),
            {"uid": uid},
        )
    ).mappings().first()

    items = []
    for r in rows:
        items.append({
            "id": str(r["id"]),
            "document_type": r["document_type"],
            "document_category": r.get("document_category", "civil"),
            "title": r.get("title") or _DOC_LABELS.get(r["document_type"], r["document_type"]),
            "generated_text": "",
            "preview_text": r.get("preview_text") or "",
            "ai_model": r.get("ai_model"),
            "used_ai": bool(r.get("used_ai", True)),
            "has_docx_export": bool(r.get("has_docx_export", False)),
            "has_pdf_export": bool(r.get("has_pdf_export", False)),
            "last_exported_at": r.get("last_exported_at"),
            "e_court_ready": bool(r.get("e_court_ready", False)),
            "filing_blockers": r.get("filing_blockers") or [],
            "case_id": str(r["case_id"]) if r.get("case_id") else None,
            "created_at": r["created_at"].isoformat() if hasattr(r["created_at"], "isoformat") else str(r["created_at"]),
        })

    return {
        "user_id": uid,
        "total": total,
        "page": page,
        "page_size": page_size,
        "pages": max(1, (total + page_size - 1) // page_size),
        "sort_by": sort_by,
        "sort_dir": sort_dir,
        "query": query,
        "doc_type": doc_type,
        "has_docx_export": None,
        "has_pdf_export": None,
        "items": items,
        "usage": {
            "docs_used": sub_row["docs_used"] if sub_row else 0,
            "docs_limit": sub_row["docs_limit"] if sub_row else 5,
        },
    }


@app.get("/api/documents/{doc_id}")
async def get_document(
    doc_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    row = (
        await session.execute(
            text("SELECT * FROM generated_documents WHERE id = :id AND user_id = :uid LIMIT 1"),
            {"id": doc_id, "uid": str(current_user["id"])},
        )
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Document not found")
    return _doc_row_to_dict(row)


@app.patch("/api/documents/{doc_id}")
async def update_document(
    doc_id: str,
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    exists = (
        await session.execute(
            text("SELECT 1 FROM generated_documents WHERE id = :id AND user_id = :uid"), {"id": doc_id, "uid": uid}
        )
    ).first()
    if not exists:
        raise HTTPException(status_code=404, detail="Document not found")

    fields: dict = {}
    if "generated_text" in body:
        # Save current version before overwriting
        try:
            cur = (await session.execute(
                text("SELECT generated_text FROM generated_documents WHERE id = :id LIMIT 1"), {"id": doc_id}
            )).scalar()
            max_ver = (await session.execute(
                text("SELECT COALESCE(MAX(version_number), 0) FROM document_versions WHERE document_id = :did"),
                {"did": doc_id},
            )).scalar() or 0
            await session.execute(
                text("INSERT INTO document_versions (id, document_id, version_number, action, generated_text) VALUES (:id, :did, :ver, 'edit', :text)"),
                {"id": str(uuid.uuid4()), "did": doc_id, "ver": max_ver + 1, "text": cur or ""},
            )
        except Exception as ve:
            print(f"[versioning] {ve}")
        fields["generated_text"] = body["generated_text"]
        fields["preview_text"] = body["generated_text"][:200] if body["generated_text"] else ""
    if "case_id" in body:
        fields["case_id"] = body["case_id"]
    if not fields:
        return {"status": "ok", "id": doc_id, "has_docx_export": False, "has_pdf_export": False}

    sets = ", ".join(f"{k} = :{k}" for k in fields)
    fields["id"] = doc_id
    await session.execute(text(f"UPDATE generated_documents SET {sets} WHERE id = :id"), fields)
    await session.commit()
    await _audit_log(session, uid, "document_update", "document", doc_id)
    return {"status": "ok", "id": doc_id, "has_docx_export": False, "has_pdf_export": False}


@app.delete("/api/documents/{doc_id}")
async def delete_document(
    doc_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    await session.execute(
        text("DELETE FROM generated_documents WHERE id = :id AND user_id = :uid"),
        {"id": doc_id, "uid": uid},
    )
    await session.commit()
    await _audit_log(session, uid, "document_delete", "document", doc_id)
    return {"status": "deleted", "id": doc_id}


def _doc_row_to_dict(r: Any) -> dict:
    return {
        "id": str(r["id"]),
        "document_type": r["document_type"],
        "document_category": r.get("document_category", "civil"),
        "title": r.get("title") or _DOC_LABELS.get(r["document_type"], r["document_type"]),
        "generated_text": r.get("generated_text", ""),
        "preview_text": r.get("preview_text") or "",
        "ai_model": r.get("ai_model"),
        "used_ai": bool(r.get("used_ai", True)),
        "has_docx_export": bool(r.get("has_docx_export", False)),
        "has_pdf_export": bool(r.get("has_pdf_export", False)),
        "last_exported_at": r.get("last_exported_at"),
        "e_court_ready": bool(r.get("e_court_ready", False)),
        "filing_blockers": r.get("filing_blockers") or [],
        "case_id": str(r["case_id"]) if r.get("case_id") else None,
        "created_at": r["created_at"].isoformat() if hasattr(r["created_at"], "isoformat") else str(r["created_at"]),
    }


# ── Document generation ───────────────────────────────────────────────────────
class GenerateRequest(BaseModel):
    doc_type: str
    form_data: dict = {}
    tariff: str = "FREE"
    extra_prompt_context: str | None = None
    case_id: str | None = None
    mode: str | None = None
    style: str | None = None


class AsyncGenerateRequest(BaseModel):
    doc_type: str
    form_data: dict = {}
    tariff: str = "FREE"
    extra_prompt_context: str | None = None


class AsyncJobStatusResponse(BaseModel):
    job_id: str
    task_id: str
    status: str
    progress: int = 0
    message: str | None = None
    result: dict | None = None
    error: str | None = None


async def _create_async_job(
    session: AsyncSession,
    uid: str,
    task_id: str,
    job_type: str,
    payload: dict[str, Any],
) -> str:
    job_id = str(uuid.uuid4())
    await session.execute(
        text(
            """
            INSERT INTO async_jobs (id, user_id, task_id, job_type, status, payload)
            VALUES (:id, :uid, :task_id, :job_type, 'queued', :payload::jsonb)
            """
        ),
        {"id": job_id, "uid": uid, "task_id": task_id, "job_type": job_type, "payload": json.dumps(payload, ensure_ascii=False)},
    )
    await session.commit()
    return job_id


async def _dead_letter_once(
    session: AsyncSession,
    uid: str,
    task_id: str,
    job_type: str,
    error_message: str,
    meta: dict[str, Any] | None = None,
) -> None:
    exists = (
        await session.execute(
            text("SELECT 1 FROM dead_letter_jobs WHERE task_id = :task_id AND user_id = :uid LIMIT 1"),
            {"task_id": task_id, "uid": uid},
        )
    ).first()
    if exists:
        return
    await session.execute(
        text(
            """
            INSERT INTO dead_letter_jobs (id, user_id, task_id, job_type, error_message, meta)
            VALUES (:id, :uid, :task_id, :job_type, :error_message, :meta::jsonb)
            """
        ),
        {
            "id": str(uuid.uuid4()),
            "uid": uid,
            "task_id": task_id,
            "job_type": job_type,
            "error_message": error_message[:2000],
            "meta": json.dumps(meta or {}, ensure_ascii=False),
        },
    )
    await session.commit()


@app.post("/api/documents/generate")
async def generate_document(
    body: GenerateRequest,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])

    # Check quota
    sub_row = (
        await session.execute(
            text("SELECT plan, docs_used, docs_limit FROM subscriptions WHERE user_id = :uid LIMIT 1"),
            {"uid": uid},
        )
    ).mappings().first()

    docs_used = sub_row["docs_used"] if sub_row else 0
    docs_limit = sub_row["docs_limit"] if sub_row else 5

    if docs_limit is not None and docs_used >= docs_limit:
        raise HTTPException(
            status_code=402,
            detail=f"Ліміт документів вичерпано ({docs_used}/{docs_limit}). Оновіть тариф.",
        )

    # Legal Brain grounding: retrieval -> prompt context -> citations
    retrieval_query = " ".join([body.doc_type, *[str(v) for v in body.form_data.values() if v]])[:1500]
    retrieved_chunks = await _retrieve_legal_chunks(session, uid, retrieval_query, top_k=5)
    citations = _build_citations(retrieved_chunks)
    citation_context = ""
    if citations:
        blocks = []
        for c in citations:
            blocks.append(
                f"[{c['citation_id']}] {c.get('code') or ''} {c.get('locator') or ''}\n"
                f"{c.get('evidence_span') or ''}"
            )
        citation_context = "Нормативні джерела (RAG):\n" + "\n\n".join(blocks)
    merged_extra_context = "\n\n".join([x for x in [citation_context, body.extra_prompt_context] if x])

    # Generate with AI or fallback
    generated_text = await _generate_text(body.doc_type, body.form_data, merged_extra_context or None)
    quality_metrics = _grounding_quality_metrics(generated_text, citations)
    preview = generated_text[:200] if generated_text else ""
    doc_title = _DOC_LABELS.get(body.doc_type, body.doc_type.replace("_", " ").title())
    ai_model = "claude-sonnet-4-6"
    doc_id = str(uuid.uuid4())
    category = _DOC_CATEGORIES.get(body.doc_type, "civil")

    await session.execute(
        text("""
            INSERT INTO generated_documents
              (id, user_id, case_id, document_type, document_category, title,
               generated_text, preview_text, ai_model, used_ai, created_at, updated_at)
            VALUES
              (:id, :uid, :cid, :dt, :cat, :title,
               :text, :preview, :model, true, NOW(), NOW())
        """),
        {
            "id": doc_id, "uid": uid, "cid": body.case_id,
            "dt": body.doc_type, "cat": category, "title": doc_title,
            "text": generated_text, "preview": preview, "model": ai_model,
        },
    )

    # Increment usage counter
    await session.execute(
        text("""
            INSERT INTO subscriptions (user_id, plan, docs_used, docs_limit)
            VALUES (:uid, 'FREE', 1, 5)
            ON CONFLICT (user_id) DO UPDATE SET docs_used = subscriptions.docs_used + 1
        """),
        {"uid": uid},
    )
    await session.commit()

    return {
        "status": "generated",
        "id": doc_id,
        "document_type": body.doc_type,
        "document_category": category,
        "title": doc_title,
        "generated_text": generated_text,
        "preview_text": preview,
        "ai_model": ai_model,
        "used_ai": True,
        "has_docx_export": False,
        "has_pdf_export": False,
        "e_court_ready": False,
        "filing_blockers": [],
        "citations": citations,
        "quality_metrics": quality_metrics,
        "created_at": datetime.utcnow().isoformat(),
    }


@app.post("/api/jobs/generate")
async def enqueue_generate_job(
    body: AsyncGenerateRequest,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    sub_row = (
        await session.execute(
            text("SELECT docs_used, docs_limit FROM subscriptions WHERE user_id = :uid LIMIT 1"),
            {"uid": uid},
        )
    ).mappings().first()
    docs_used = sub_row["docs_used"] if sub_row else 0
    docs_limit = sub_row["docs_limit"] if sub_row else 5
    if docs_limit is not None and docs_used >= docs_limit:
        raise HTTPException(status_code=402, detail=f"Ліміт документів вичерпано ({docs_used}/{docs_limit}).")

    async_result = generate_document_job.delay(body.doc_type, body.form_data, body.extra_prompt_context)
    job_id = await _create_async_job(
        session,
        uid,
        async_result.id,
        "generate_document",
        {
            "doc_type": body.doc_type,
            "tariff": body.tariff,
            "form_data": body.form_data,
            "extra_prompt_context": body.extra_prompt_context,
        },
    )
    return {"status": "queued", "job_id": job_id, "task_id": async_result.id}


@app.post("/api/jobs/analyze-intake")
async def enqueue_intake_job(
    file: UploadFile = File(...),
    jurisdiction: str = Form("UA"),
    mode: str = Form("standard"),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    file_bytes = await file.read()
    payload = {
        "file_b64": base64.b64encode(file_bytes).decode("utf-8"),
        "file_name": file.filename or "document",
        "jurisdiction": jurisdiction,
        "mode": mode,
    }
    async_result = analyze_intake_job.delay(payload)
    job_id = await _create_async_job(session, uid, async_result.id, "analyze_intake", {"file_name": payload["file_name"]})
    return {"status": "queued", "job_id": job_id, "task_id": async_result.id}


@app.get("/api/jobs/{job_id}", response_model=AsyncJobStatusResponse)
async def get_async_job_status(
    job_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    row = (
        await session.execute(
            text("SELECT * FROM async_jobs WHERE id = :id AND user_id = :uid LIMIT 1"),
            {"id": job_id, "uid": uid},
        )
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Job not found")

    task_id = row["task_id"]
    result = celery_app.AsyncResult(task_id)
    state = (result.state or "PENDING").upper()
    meta = result.info if isinstance(result.info, dict) else {}
    progress = int(meta.get("progress", 0))
    message = meta.get("message")
    error = None
    result_payload = None
    status_value = "queued"

    if state in {"PENDING"}:
        status_value = "queued"
    elif state in {"STARTED", "PROGRESS", "RETRY"}:
        status_value = "running"
        progress = max(progress, 10)
    elif state == "SUCCESS":
        status_value = "success"
        progress = 100
        result_payload = result.result if isinstance(result.result, dict) else {"raw": result.result}
    else:
        status_value = "failed"
        progress = 100
        error = str(result.result)[:2000]
        await _dead_letter_once(
            session,
            uid,
            task_id,
            row["job_type"],
            error,
            {"state": state},
        )

    # Persist job status + result snapshot for auditability.
    await session.execute(
        text(
            """
            UPDATE async_jobs
            SET status = :status,
                retries = :retries,
                result = COALESCE(:result::jsonb, result),
                error_message = COALESCE(:error_message, error_message),
                updated_at = NOW()
            WHERE id = :id AND user_id = :uid
            """
        ),
        {
            "status": status_value,
            "retries": int(getattr(result, "retries", 0) or 0),
            "result": json.dumps(result_payload, ensure_ascii=False) if result_payload is not None else None,
            "error_message": error,
            "id": job_id,
            "uid": uid,
        },
    )

    # Finalize generated document once per job.
    if status_value == "success" and row["job_type"] == "generate_document":
        existing = (
            await session.execute(
                text(
                    """
                    SELECT id FROM generated_documents
                    WHERE user_id = :uid AND title = :title AND created_at > NOW() - INTERVAL '10 minutes'
                    ORDER BY created_at DESC LIMIT 1
                    """
                ),
                {"uid": uid, "title": (result_payload or {}).get("title", "")},
            )
        ).mappings().first()
        if not existing:
            doc_id = str(uuid.uuid4())
            payload = result_payload or {}
            category = _DOC_CATEGORIES.get(payload.get("document_type", "pozov_do_sudu"), "civil")
            # Legal Brain post-generation grounding for async jobs.
            job_payload = row.get("payload") or {}
            if isinstance(job_payload, str):
                try:
                    job_payload = json.loads(job_payload)
                except Exception:
                    job_payload = {}
            query = " ".join(
                [
                    str(payload.get("document_type", "")),
                    *[str(v) for v in (job_payload.get("form_data") or {}).values() if v],
                ]
            )[:1500]
            retrieved_chunks = await _retrieve_legal_chunks(session, uid, query, top_k=5)
            citations = _build_citations(retrieved_chunks)
            payload["citations"] = citations
            payload["quality_metrics"] = _grounding_quality_metrics(payload.get("generated_text", ""), citations)
            await session.execute(
                text(
                    """
                    INSERT INTO generated_documents
                      (id, user_id, document_type, document_category, title,
                       generated_text, preview_text, ai_model, used_ai, created_at, updated_at)
                    VALUES
                      (:id, :uid, :dt, :cat, :title, :text, :preview, :model, :used_ai, NOW(), NOW())
                    """
                ),
                {
                    "id": doc_id,
                    "uid": uid,
                    "dt": payload.get("document_type", "pozov_do_sudu"),
                    "cat": category,
                    "title": payload.get("title", "Згенерований документ"),
                    "text": payload.get("generated_text", ""),
                    "preview": payload.get("preview_text", ""),
                    "model": payload.get("ai_model", "claude-sonnet-4-6"),
                    "used_ai": bool(payload.get("used_ai", True)),
                },
            )
            await session.execute(
                text(
                    """
                    INSERT INTO subscriptions (user_id, plan, docs_used, docs_limit)
                    VALUES (:uid, 'FREE', 1, 5)
                    ON CONFLICT (user_id) DO UPDATE SET docs_used = subscriptions.docs_used + 1
                    """
                ),
                {"uid": uid},
            )
            result_payload["document_id"] = doc_id

    await session.commit()
    return {
        "job_id": job_id,
        "task_id": task_id,
        "status": status_value,
        "progress": progress,
        "message": message,
        "result": result_payload,
        "error": error,
    }


async def _generate_text(doc_type: str, form_data: dict, extra_context: str | None) -> str:
    """Generate document text via Anthropic API, falling back to template."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        return _template_document(doc_type, form_data)

    try:
        import anthropic
        client = anthropic.Anthropic(api_key=api_key)

        label = _DOC_LABELS.get(doc_type, doc_type)
        fields_text = "\n".join(f"  {k}: {v}" for k, v in form_data.items() if v)
        extra = f"\n\nДодатковий контекст: {extra_context}" if extra_context else ""

        prompt = f"""Ти — досвідчений юрист. Склади юридичний документ: {label}.

Дані для документа:
{fields_text}{extra}

Вимоги:
- Документ має бути написаний українською мовою
- Дотримуйся юридичних стандартів і формату
- Включи всі необхідні реквізити
- Документ повинен бути готовим до використання

Надай тільки текст документа без пояснень."""

        message = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=4096,
            messages=[{"role": "user", "content": prompt}],
        )
        return message.content[0].text
    except Exception as e:
        print(f"[generate] AI error: {e}")
        return _template_document(doc_type, form_data)


def _template_document(doc_type: str, form_data: dict) -> str:
    """Simple template fallback when AI is unavailable."""
    label = _DOC_LABELS.get(doc_type, doc_type.replace("_", " ").title())
    plaintiff = form_data.get("plaintiff_name") or form_data.get("party_name") or form_data.get("name") or "[Ім'я]"
    defendant = form_data.get("defendant_name") or form_data.get("counterparty") or "[Відповідач]"
    description = form_data.get("description") or form_data.get("claim_description") or form_data.get("subject") or "[Опис]"
    amount = form_data.get("amount") or form_data.get("claim_amount") or ""
    date = datetime.utcnow().strftime("%d.%m.%Y")

    return f"""{label.upper()}

Дата складення: {date}

СТОРОНИ:
Позивач / Сторона 1: {plaintiff}
Відповідач / Сторона 2: {defendant}

СУТЬ ЗВЕРНЕННЯ:
{description}
{f"Сума вимоги: {amount} грн." if amount else ""}

На підставі вищевикладеного прошу задовольнити вимоги в повному обсязі.

Дата: {date}
Підпис: _______________
"""


# ============================================================================
# CASES
# ============================================================================
@app.get("/api/cases")
async def get_cases(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    rows = (
        await session.execute(
            text("""
                SELECT id, user_id, title, description, case_number, status, created_at, updated_at
                FROM cases WHERE user_id = :uid ORDER BY created_at DESC
            """),
            {"uid": uid},
        )
    ).mappings().all()
    return [_case_row(r) for r in rows]


@app.post("/api/cases", status_code=201)
async def create_case(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    case_id = str(uuid.uuid4())
    now = datetime.utcnow()
    await session.execute(
        text("""
            INSERT INTO cases (id, user_id, title, description, case_number, status, created_at, updated_at)
            VALUES (:id, :uid, :title, :desc, :num, :status, :now, :now)
        """),
        {
            "id": case_id, "uid": uid,
            "title": body.get("title", "Без назви"),
            "desc": body.get("description"),
            "num": body.get("case_number"),
            "status": body.get("status", "open"),
            "now": now,
        },
    )
    await session.commit()
    row = (
        await session.execute(
            text("SELECT * FROM cases WHERE id = :id LIMIT 1"), {"id": case_id}
        )
    ).mappings().first()
    return _case_row(row)


@app.get("/api/cases/{case_id}")
async def get_case(
    case_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    row = (
        await session.execute(
            text("SELECT * FROM cases WHERE id = :id AND user_id = :uid LIMIT 1"),
            {"id": case_id, "uid": str(current_user["id"])},
        )
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Case not found")

    docs = (
        await session.execute(
            text("SELECT id, document_type, document_category, created_at FROM generated_documents WHERE case_id = :cid ORDER BY created_at DESC"),
            {"cid": case_id},
        )
    ).mappings().all()

    result = _case_row(row)
    result["documents"] = [
        {"id": str(d["id"]), "document_type": d["document_type"],
         "document_category": d.get("document_category", "civil"),
         "created_at": d["created_at"].isoformat() if hasattr(d["created_at"], "isoformat") else str(d["created_at"])}
        for d in docs
    ]
    result["forum_posts"] = []
    result["case_law_items"] = []
    return result


@app.patch("/api/cases/{case_id}")
async def update_case(
    case_id: str,
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    allowed = {"title", "description", "case_number", "status"}
    fields = {k: v for k, v in body.items() if k in allowed}
    if fields:
        fields["updated_at"] = datetime.utcnow()
        sets = ", ".join(f"{k} = :{k}" for k in fields)
        fields["id"] = case_id
        fields["uid"] = uid
        await session.execute(
            text(f"UPDATE cases SET {sets} WHERE id = :id AND user_id = :uid"), fields
        )
        await session.commit()
    row = (
        await session.execute(
            text("SELECT * FROM cases WHERE id = :id AND user_id = :uid LIMIT 1"),
            {"id": case_id, "uid": uid},
        )
    ).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Case not found")
    return _case_row(row)


@app.delete("/api/cases/{case_id}")
async def delete_case(
    case_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    await session.execute(
        text("DELETE FROM cases WHERE id = :id AND user_id = :uid"),
        {"id": case_id, "uid": str(current_user["id"])},
    )
    await session.commit()
    return {"status": "deleted"}


def _case_row(r: Any) -> dict:
    return {
        "id": str(r["id"]),
        "user_id": str(r["user_id"]),
        "title": r["title"],
        "description": r.get("description"),
        "case_number": r.get("case_number"),
        "status": r.get("status", "open"),
        "created_at": r["created_at"].isoformat() if hasattr(r["created_at"], "isoformat") else str(r["created_at"]),
        "updated_at": r["updated_at"].isoformat() if hasattr(r.get("updated_at"), "isoformat") else str(r.get("updated_at", r["created_at"])),
    }


# ============================================================================
# GLOBAL SEARCH
# ============================================================================
@app.get("/api/dashboard/search")
async def global_search(
    q: str = Query(..., min_length=2),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    pattern = f"%{q}%"

    cases = (
        await session.execute(
            text("SELECT id, title, case_number FROM cases WHERE user_id = :uid AND title ILIKE :q LIMIT 5"),
            {"uid": uid, "q": pattern},
        )
    ).mappings().all()

    docs = (
        await session.execute(
            text("SELECT id, document_type, preview_text FROM generated_documents WHERE user_id = :uid AND (title ILIKE :q OR preview_text ILIKE :q) LIMIT 5"),
            {"uid": uid, "q": pattern},
        )
    ).mappings().all()

    forum_rows = []
    try:
        forum_rows = (await session.execute(
            text("SELECT id, title FROM forum_posts WHERE (user_id = :uid OR 1=1) AND title ILIKE :q LIMIT 5"),
            {"uid": uid, "q": pattern},
        )).mappings().all()
    except Exception:
        pass

    return {
        "cases": [{"id": str(c["id"]), "title": c["title"], "number": c.get("case_number")} for c in cases],
        "documents": [{"id": str(d["id"]), "type": d["document_type"], "preview": d.get("preview_text", "")} for d in docs],
        "forum": [{"id": str(f["id"]), "title": f["title"]} for f in forum_rows],
    }


def _required_processual_sections(doc_type: str) -> list[str]:
    if doc_type in {"pozov_do_sudu", "pozov_trudovyi"}:
        return ["СУД", "ПОЗИВАЧ", "ВІДПОВІДАЧ", "ОБСТАВИНИ", "ПРАВОВЕ ОБҐРУНТУВАННЯ", "ПРОШУ"]
    if doc_type == "appeal_complaint":
        return ["АПЕЛЯЦІЙНА СКАРГА", "СУД ПЕРШОЇ ІНСТАНЦІЇ", "ПІДСТАВИ ОСКАРЖЕННЯ", "ПРОШУ"]
    if doc_type == "pretenziya":
        return ["ПРЕТЕНЗІЯ", "ОПИС ПОРУШЕННЯ", "ВИМАГАЮ"]
    return []


def _find_processual_blockers(doc_type: str, text_value: str) -> list[str]:
    text_u = (text_value or "").upper()
    blockers: list[str] = []
    for section in _required_processual_sections(doc_type):
        if section not in text_u:
            blockers.append(f"missing_section:{section}")
    return blockers


def _apply_rule_based_processual_repair(doc_type: str, text_value: str) -> str:
    repaired = (text_value or "").strip()
    if not repaired:
        repaired = _template_document(doc_type, {})
    for section in _required_processual_sections(doc_type):
        if section not in repaired.upper():
            repaired += f"\n\n{section}\n[Доповнити розділ]"
    return repaired


async def _repair_processual_text(doc_type: str, text_value: str) -> tuple[str, list[str], list[str]]:
    blockers_before = _find_processual_blockers(doc_type, text_value)
    repaired_text = _apply_rule_based_processual_repair(doc_type, text_value)
    blockers_after_rules = _find_processual_blockers(doc_type, repaired_text)

    if blockers_after_rules and os.getenv("ANTHROPIC_API_KEY"):
        prompt = (
            "Виправ процесуальний документ українською мовою. "
            "Додай відсутні розділи без зміни суті.\n\n"
            f"Тип документа: {doc_type}\n"
            f"Обов'язкові розділи: {', '.join(_required_processual_sections(doc_type))}\n"
            f"Поточний текст:\n{text_value or ''}\n"
        )
        try:
            import anthropic

            client = anthropic.Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
            message = client.messages.create(
                model=os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6"),
                max_tokens=3000,
                messages=[{"role": "user", "content": prompt}],
            )
            ai_text = message.content[0].text if message and message.content else repaired_text
            if ai_text:
                repaired_text = ai_text
        except Exception:
            pass

    blockers_after = _find_processual_blockers(doc_type, repaired_text)
    return repaired_text, blockers_before, blockers_after


async def _sync_submission_with_adapter(submission: dict[str, Any]) -> tuple[bool, str]:
    provider = (submission.get("provider") or "manual").lower()
    tracking_url = (submission.get("tracking_url") or "").strip()

    if provider == "manual":
        return False, "manual_provider_no_live_adapter"

    if tracking_url:
        try:
            async with httpx.AsyncClient(timeout=8.0) as client:
                resp = await client.get(tracking_url)
            if 200 <= resp.status_code < 300:
                return True, "tracking_url_ok"
            return False, f"tracking_url_status_{resp.status_code}"
        except Exception as exc:
            return False, f"tracking_url_error:{str(exc)[:120]}"

    if provider == "opendatabot":
        if not os.getenv("OPENDATABOT_API_KEY"):
            return False, "missing_opendatabot_api_key"
        return True, "provider_ready"

    return False, f"provider_not_supported:{provider}"


# ============================================================================
# STUB ENDPOINTS (return empty/default so frontend doesn't break)
# ============================================================================
@app.get("/api/documents/form-schema/{doc_type}")
async def get_form_schema(doc_type: str):
    schema = _FORM_SCHEMA_CATALOG.get(doc_type)
    if not schema:
        raise HTTPException(status_code=404, detail=f"Невідомий тип документа: {doc_type}")
    return {"doc_type": doc_type, **schema}


@app.get("/api/billing/invoices")
async def get_invoices(current_user: dict = Depends(get_current_user)):
    return {"items": []}


@app.get("/api/team/users")
async def get_team_users(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    rows = (await session.execute(
        text("SELECT id, email, full_name, role, company FROM users WHERE id = :uid OR role IN ('admin','owner') LIMIT 50"),
        {"uid": uid},
    )).mappings().all()
    return {"items": [dict(r) for r in rows]}


# ============================================================================
# DEADLINES — повний CRUD з БД
# ============================================================================

@app.get("/api/deadlines")
async def get_deadlines(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    rows = (await session.execute(
        text("SELECT * FROM deadlines WHERE user_id = :uid ORDER BY end_date ASC NULLS LAST, created_at DESC LIMIT 200"),
        {"uid": uid},
    )).mappings().all()
    return {"total": len(rows), "items": [dict(r) for r in rows]}


@app.post("/api/deadlines", status_code=201)
async def create_deadline(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    did = str(uuid.uuid4())
    await session.execute(
        text("""
            INSERT INTO deadlines (id, user_id, case_id, document_id, title, deadline_type,
                start_date, end_date, reminder_sent, notes)
            VALUES (:id, :uid, :case_id, :doc_id, :title, :dtype,
                :start, :end, false, :notes)
        """),
        {
            "id": did, "uid": uid,
            "case_id": body.get("case_id"),
            "doc_id": body.get("document_id"),
            "title": body.get("title", "Строк"),
            "dtype": body.get("deadline_type"),
            "start": body.get("start_date"),
            "end": body.get("end_date"),
            "notes": body.get("notes"),
        },
    )
    await session.commit()
    row = (await session.execute(
        text("SELECT * FROM deadlines WHERE id = :id LIMIT 1"), {"id": did}
    )).mappings().first()
    return dict(row) if row else {"id": did, "user_id": uid, **body}


@app.patch("/api/deadlines/{deadline_id}")
async def update_deadline(
    deadline_id: str,
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    allowed = {"title", "deadline_type", "start_date", "end_date", "notes", "reminder_sent"}
    fields = {k: v for k, v in body.items() if k in allowed}
    if fields:
        sets = ", ".join(f"{k} = :{k}" for k in fields)
        fields.update({"id": deadline_id, "uid": uid})
        await session.execute(
            text(f"UPDATE deadlines SET {sets} WHERE id = :id AND user_id = :uid"), fields
        )
        await session.commit()
    return {"status": "ok"}


@app.delete("/api/deadlines/{deadline_id}")
async def delete_deadline(
    deadline_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    await session.execute(
        text("DELETE FROM deadlines WHERE id = :id AND user_id = :uid"),
        {"id": deadline_id, "uid": uid},
    )
    await session.commit()
    return {"status": "ok"}


@app.post("/api/deadlines/notify-due")
async def notify_due_deadlines(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    horizon_hours = int(body.get("horizon_hours", 24))
    limit = max(1, min(int(body.get("limit", 100)), 500))
    rows = (await session.execute(
        text("""
            SELECT id, title, end_date
            FROM deadlines
            WHERE user_id = :uid
              AND reminder_sent = false
              AND end_date IS NOT NULL
              AND end_date <= NOW() + (:horizon || ' hours')::interval
            ORDER BY end_date ASC
            LIMIT :lim
        """),
        {"uid": uid, "horizon": horizon_hours, "lim": limit},
    )).mappings().all()

    notifications = []
    for row in rows:
        nid = str(uuid.uuid4())
        message = f"Наближається дедлайн: {row.get('title')} ({row.get('end_date')})"
        await session.execute(
            text("""
                INSERT INTO deadline_notifications (id, user_id, deadline_id, channel, message, scheduled_for, sent_at, status)
                VALUES (:id, :uid, :did, 'in_app', :msg, NOW(), NOW(), 'sent')
            """),
            {"id": nid, "uid": uid, "did": row["id"], "msg": message},
        )
        await session.execute(
            text("UPDATE deadlines SET reminder_sent = true WHERE id = :id AND user_id = :uid"),
            {"id": row["id"], "uid": uid},
        )
        notifications.append({"id": nid, "deadline_id": str(row["id"]), "message": message})

    await session.commit()
    return {"processed": len(notifications), "items": notifications}


@app.get("/api/deadlines/notifications")
async def get_deadline_notifications(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text("""
            SELECT * FROM deadline_notifications
            WHERE user_id = :uid
            ORDER BY created_at DESC
            LIMIT :lim OFFSET :off
        """),
        {"uid": uid, "lim": page_size, "off": offset},
    )).mappings().all()
    total = (await session.execute(
        text("SELECT COUNT(*) FROM deadline_notifications WHERE user_id = :uid"),
        {"uid": uid},
    )).scalar() or 0
    return {"total": int(total), "page": page, "page_size": page_size, "items": [dict(r) for r in rows]}


# ============================================================================
# CASE-LAW — пошук в локальній БД + OpenDataBot
# ============================================================================

@app.get("/api/case-law/search")
async def search_case_law(
    q: str = Query(""),
    court_form: str | None = Query(None),
    date_from: str | None = Query(None),
    date_to: str | None = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    sort_by: str = Query("decision_date"),
    sort_dir: str = Query("desc"),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    conditions = ["user_id = :uid"]
    params: dict = {"uid": uid}
    if q.strip():
        conditions.append("(summary ILIKE :q OR case_number ILIKE :q OR court_name ILIKE :q)")
        params["q"] = f"%{q.strip()}%"
    if date_from:
        conditions.append("decision_date >= :date_from")
        params["date_from"] = date_from
    if date_to:
        conditions.append("decision_date <= :date_to")
        params["date_to"] = date_to
    where = " AND ".join(conditions)
    safe_sort = sort_by if sort_by in ("decision_date", "created_at", "relevance_score") else "decision_date"
    safe_dir = "DESC" if sort_dir == "desc" else "ASC"
    count_row = (await session.execute(
        text(f"SELECT COUNT(*) FROM case_law_items WHERE {where}"), params
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text(f"SELECT * FROM case_law_items WHERE {where} ORDER BY {safe_sort} {safe_dir} LIMIT :lim OFFSET :off"),
        {**params, "lim": page_size, "off": offset},
    )).mappings().all()
    return {
        "total": count_row, "page": page, "page_size": page_size,
        "pages": max(1, -(-int(count_row) // page_size)),
        "sort_by": sort_by, "sort_dir": sort_dir,
        "items": [dict(r) for r in rows],
    }


@app.get("/api/case-law/digest")
async def get_case_law_digest(
    q: str | None = Query(None),
    page: int = Query(1, ge=1),
    page_size: int = Query(20),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    params: dict = {"uid": uid}
    conditions = ["user_id = :uid"]
    if q:
        conditions.append("(title ILIKE :q OR summary ILIKE :q)")
        params["q"] = f"%{q}%"
    where = " AND ".join(conditions)
    count_row = (await session.execute(
        text(f"SELECT COUNT(*) FROM case_law_digest WHERE {where}"), params
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text(f"SELECT * FROM case_law_digest WHERE {where} ORDER BY created_at DESC LIMIT :lim OFFSET :off"),
        {**params, "lim": page_size, "off": offset},
    )).mappings().all()
    return {
        "total": count_row, "page": page, "page_size": page_size,
        "pages": max(1, -(-int(count_row) // page_size)),
        "items": [dict(r) for r in rows],
    }


@app.get("/api/case-law/digest/history")
async def get_digest_history(
    page: int = Query(1, ge=1),
    page_size: int = Query(20),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    count_row = (await session.execute(
        text("SELECT COUNT(*) FROM case_law_digest WHERE user_id = :uid"), {"uid": uid}
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text("SELECT * FROM case_law_digest WHERE user_id = :uid ORDER BY created_at DESC LIMIT :lim OFFSET :off"),
        {"uid": uid, "lim": page_size, "off": offset},
    )).mappings().all()
    return {"total": count_row, "page": page, "page_size": page_size,
            "pages": max(1, -(-int(count_row) // page_size)), "items": [dict(r) for r in rows]}


@app.get("/api/case-law/digest/history/{digest_id}")
async def get_digest_detail(
    digest_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    row = (await session.execute(
        text("SELECT * FROM case_law_digest WHERE id = :id LIMIT 1"), {"id": digest_id}
    )).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Дайджест не знайдено")
    return dict(row)


@app.get("/api/case-law/sync/status")
async def get_caselaw_sync_status(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    total = (await session.execute(
        text("SELECT COUNT(*) FROM case_law_items WHERE user_id = :uid"), {"uid": uid}
    )).scalar() or 0
    latest = (await session.execute(
        text("SELECT MAX(decision_date) FROM case_law_items WHERE user_id = :uid"), {"uid": uid}
    )).scalar()
    oldest = (await session.execute(
        text("SELECT MIN(decision_date) FROM case_law_items WHERE user_id = :uid"), {"uid": uid}
    )).scalar()
    return {
        "total_records": total, "sources": {"opendatabot": total},
        "latest_decision_date": str(latest) if latest else None,
        "oldest_decision_date": str(oldest) if oldest else None,
        "last_sync_at": None, "last_sync_action": None, "last_sync_query": None,
        "last_sync_limit": None, "last_sync_created": None, "last_sync_updated": None,
        "last_sync_total": None, "last_sync_sources": [], "last_sync_seed_fallback_used": None,
    }


# ============================================================================
# KNOWLEDGE BASE — виправляємо GET без слешу
# ============================================================================

@app.get("/api/knowledge-base")
async def get_knowledge_base(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    try:
        rows = (await session.execute(
            text("SELECT * FROM knowledge_entries WHERE user_id = :uid ORDER BY created_at DESC LIMIT 100"),
            {"uid": uid},
        )).mappings().all()
        return {"items": [dict(r) for r in rows]}
    except Exception:
        return {"items": []}


# ============================================================================
# REGISTRY WATCH — повний CRUD з БД
# ============================================================================

@app.get("/api/registries/watch")
async def get_registry_watch(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    try:
        rows = (await session.execute(
            text("SELECT * FROM registry_watch_items WHERE user_id = :uid ORDER BY created_at DESC LIMIT 100"),
            {"uid": uid},
        )).mappings().all()
        total = len(rows)
        return {"total": total, "page": 1, "page_size": 100, "pages": 1, "items": [dict(r) for r in rows]}
    except Exception:
        return {"total": 0, "page": 1, "page_size": 100, "pages": 1, "items": []}


# ============================================================================
# MONITORING STATUS
# ============================================================================

@app.get("/api/monitoring/status")
async def get_monitoring_status(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    try:
        total = (await session.execute(
            text("SELECT COUNT(*) FROM registry_watch_items WHERE user_id = :uid"), {"uid": uid}
        )).scalar() or 0
        active = (await session.execute(
            text("SELECT COUNT(*) FROM registry_watch_items WHERE user_id = :uid AND status = 'active'"), {"uid": uid}
        )).scalar() or 0
        events_24h = (await session.execute(
            text("SELECT COUNT(*) FROM registry_events WHERE user_id = :uid AND created_at > NOW() - INTERVAL '24 hours'"), {"uid": uid}
        )).scalar() or 0
    except Exception:
        total, active, events_24h = 0, 0, 0
    return {
        "total_watch_items": total, "active_watch_items": active,
        "due_watch_items": 0, "warning_watch_items": 0,
        "state_changed_events_24h": events_24h,
        "last_event_at": None, "by_status": {"active": active},
        "status": "active" if active > 0 else "inactive",
    }


# ============================================================================
# REPORTS — генерація з реальних даних
# ============================================================================

@app.get("/api/reports")
async def get_reports(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    reports = []
    try:
        docs = (await session.execute(
            text("SELECT COUNT(*) FROM generated_documents WHERE user_id = :uid"), {"uid": uid}
        )).scalar() or 0
        cases = (await session.execute(
            text("SELECT COUNT(*) FROM cases WHERE user_id = :uid"), {"uid": uid}
        )).scalar() or 0
        analyses = (await session.execute(
            text("SELECT COUNT(*) FROM document_intakes WHERE user_id = :uid"), {"uid": uid}
        )).scalar() or 0
        now = datetime.utcnow().isoformat()
        reports = [
            {"id": "summary", "title": "Загальний звіт активності", "type": "summary",
             "data": {"documents_generated": docs, "cases_total": cases, "analyses_done": analyses},
             "created_at": now},
        ]
    except Exception:
        pass
    return {"items": reports}


# ============================================================================
# CALCULATIONS — реальна логіка + зберігання в БД
# ============================================================================

@app.get("/api/calculations/history")
@app.get("/api/calculate/history")
async def get_calculation_history(
    page: int = Query(1, ge=1),
    page_size: int = Query(20),
    calculation_type: str | None = Query(None),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    params: dict = {"uid": uid}
    conditions = ["user_id = :uid"]
    if calculation_type:
        conditions.append("calculation_type = :ctype")
        params["ctype"] = calculation_type
    where = " AND ".join(conditions)
    count_row = (await session.execute(
        text(f"SELECT COUNT(*) FROM calculations WHERE {where}"), params
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text(f"SELECT * FROM calculations WHERE {where} ORDER BY created_at DESC LIMIT :lim OFFSET :off"),
        {**params, "lim": page_size, "off": offset},
    )).mappings().all()
    items = []
    for r in rows:
        d = dict(r)
        for f in ("input_payload", "output_payload"):
            if isinstance(d.get(f), str):
                try:
                    d[f] = json.loads(d[f])
                except Exception:
                    d[f] = {}
        items.append(d)
    return {"total": count_row, "page": page, "page_size": page_size,
            "pages": max(1, -(-int(count_row) // page_size)), "items": items}


@app.get("/api/calculate/{calc_id}")
async def get_calculation_detail(
    calc_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    row = (await session.execute(
        text("SELECT * FROM calculations WHERE id = :id AND user_id = :uid LIMIT 1"),
        {"id": calc_id, "uid": uid},
    )).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Розрахунок не знайдено")
    d = dict(row)
    for f in ("input_payload", "output_payload"):
        if isinstance(d.get(f), str):
            try:
                d[f] = json.loads(d[f])
            except Exception:
                d[f] = {}
    return {"item": d}


@app.post("/api/calculations/full-claim")
@app.post("/api/calculate/full")
async def calculate_full_claim(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    principal = float(body.get("principal_amount", body.get("principal", 0)))
    rate = float(body.get("annual_rate", body.get("rate", 0.3)))
    days = int(body.get("days_overdue", body.get("days", 0)))
    court_fee_rate = float(body.get("court_fee_rate", 0.01))

    penalty = round(principal * rate / 365 * days, 2)
    court_fee = round(max(principal * court_fee_rate, 640), 2)
    total_claim = round(principal + penalty, 2)
    total_with_fee = round(total_claim + court_fee, 2)

    from datetime import date, timedelta  # noqa: PLC0415
    today = date.today()
    process_deadline = (today + timedelta(days=30)).isoformat()
    limitation_deadline = (today + timedelta(days=365 * 3)).isoformat()

    result = {
        "court_fee_uah": court_fee,
        "penalty_uah": penalty,
        "process_deadline": process_deadline,
        "limitation_deadline": limitation_deadline,
        "total_claim_uah": total_claim,
        "total_with_fee_uah": total_with_fee,
    }

    cid = str(uuid.uuid4())
    try:
        await session.execute(
            text("""
                INSERT INTO calculations (id, user_id, calculation_type, title, input_payload, output_payload)
                VALUES (:id, :uid, 'full_claim', :title, :inp::jsonb, :out::jsonb)
            """),
            {
                "id": cid, "uid": uid,
                "title": f"Претензія {principal:,.0f} грн × {days} днів",
                "inp": json.dumps(body, ensure_ascii=False),
                "out": json.dumps(result, ensure_ascii=False),
            },
        )
        await session.commit()
        saved = True
    except Exception as e:
        print(f"[calculations] DB error: {e}")
        saved = False
        cid = None

    return {"status": "ok", "result": result, "saved": saved,
            "calculation_id": cid, "created_at": datetime.utcnow().isoformat()}


# ============================================================================
# AUDIT LOG — реальна таблиця
# ============================================================================

async def _audit_log(session: AsyncSession, user_id: str, action: str, entity_type: str | None = None,
                     entity_id: str | None = None, metadata: dict | None = None):
    """Записати подію в audit_logs (best-effort)."""
    try:
        prev_hash = (await session.execute(
            text("SELECT integrity_hash FROM audit_logs WHERE user_id = :uid ORDER BY created_at DESC LIMIT 1"),
            {"uid": user_id},
        )).scalar()
        payload = f"{user_id}:{action}:{entity_type}:{entity_id}:{datetime.utcnow().isoformat()}"
        curr_hash = hmac.new(_SECRET.encode(), payload.encode(), hashlib.sha256).hexdigest()
        await session.execute(
            text("""
                INSERT INTO audit_logs (id, user_id, action, entity_type, entity_id, metadata,
                    integrity_scope, integrity_prev_hash, integrity_hash)
                VALUES (:id, :uid, :action, :etype, :eid, :meta::jsonb, 'user', :prev, :curr)
            """),
            {
                "id": str(uuid.uuid4()), "uid": user_id, "action": action,
                "etype": entity_type, "eid": entity_id,
                "meta": json.dumps(metadata or {}, ensure_ascii=False),
                "prev": prev_hash, "curr": curr_hash,
            },
        )
    except Exception as e:
        print(f"[audit_log] error: {e}")


@app.get("/api/audit/history")
async def get_audit_history(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    action: str | None = Query(None),
    entity_type: str | None = Query(None),
    query: str | None = Query(None),
    sort_dir: str = Query("desc"),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    conditions = ["user_id = :uid"]
    params: dict = {"uid": uid}
    if action:
        conditions.append("action = :action")
        params["action"] = action
    if entity_type:
        conditions.append("entity_type = :etype")
        params["etype"] = entity_type
    if query:
        conditions.append("(action ILIKE :q OR entity_type ILIKE :q)")
        params["q"] = f"%{query}%"
    where = " AND ".join(conditions)
    safe_dir = "DESC" if sort_dir == "desc" else "ASC"
    count_row = (await session.execute(
        text(f"SELECT COUNT(*) FROM audit_logs WHERE {where}"), params
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text(f"SELECT * FROM audit_logs WHERE {where} ORDER BY created_at {safe_dir} LIMIT :lim OFFSET :off"),
        {**params, "lim": page_size, "off": offset},
    )).mappings().all()
    items = []
    for r in rows:
        d = dict(r)
        if isinstance(d.get("metadata"), str):
            try:
                d["metadata"] = json.loads(d["metadata"])
            except Exception:
                d["metadata"] = {}
        items.append(d)
    return {
        "user_id": uid, "total": count_row, "page": page, "page_size": page_size,
        "pages": max(1, -(-int(count_row) // page_size)),
        "action": action, "entity_type": entity_type, "query": query, "items": items,
    }


@app.get("/api/audit/integrity")
async def get_audit_integrity(
    max_rows: int = Query(100),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    rows = (await session.execute(
        text("SELECT * FROM audit_logs WHERE user_id = :uid ORDER BY created_at ASC LIMIT :lim"),
        {"uid": uid, "lim": max_rows},
    )).mappings().all()
    issues = []
    for i, r in enumerate(rows):
        if i > 0 and r.get("integrity_prev_hash") != rows[i - 1].get("integrity_hash"):
            issues.append({
                "row_id": str(r["id"]), "created_at": str(r.get("created_at")),
                "code": "CHAIN_BROKEN", "message": "Порушено ланцюг хешів",
            })
    return {"scope": "user", "total_checked": len(rows), "issues": issues, "ok": len(issues) == 0}


# ============================================================================
# FORUM  (/forum/posts — без /api/ префіксу, як у frontend)
# ============================================================================

@app.get("/forum/posts")
async def get_forum_posts(
    case_id: str | None = Query(None),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    params: dict = {}
    conditions = []
    if case_id:
        conditions.append("f.case_id = :case_id")
        params["case_id"] = case_id
    where = ("WHERE " + " AND ".join(conditions)) if conditions else ""
    rows = (await session.execute(
        text(f"""
            SELECT f.*, u.full_name as user_name,
                (SELECT COUNT(*) FROM forum_comments c WHERE c.post_id = f.id) as comment_count
            FROM forum_posts f
            LEFT JOIN users u ON u.id = f.user_id
            {where}
            ORDER BY f.created_at DESC LIMIT 50
        """), params,
    )).mappings().all()
    return [dict(r) for r in rows]


@app.post("/forum/posts", status_code=201)
async def create_forum_post(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    pid = str(uuid.uuid4())
    await session.execute(
        text("""
            INSERT INTO forum_posts (id, user_id, title, content, category, case_id)
            VALUES (:id, :uid, :title, :content, :category, :case_id)
        """),
        {"id": pid, "uid": uid, "title": body.get("title", ""),
         "content": body.get("content", ""), "category": body.get("category"),
         "case_id": body.get("case_id")},
    )
    await session.commit()
    return {"id": pid, "user_id": uid, "user_name": current_user.get("full_name"),
            "title": body.get("title", ""), "content": body.get("content", ""),
            "category": body.get("category"), "case_id": body.get("case_id"),
            "created_at": datetime.utcnow().isoformat(), "comment_count": 0}


@app.get("/forum/posts/{post_id}")
async def get_forum_post(
    post_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    row = (await session.execute(
        text("""
            SELECT f.*, u.full_name as user_name
            FROM forum_posts f LEFT JOIN users u ON u.id = f.user_id
            WHERE f.id = :id LIMIT 1
        """), {"id": post_id},
    )).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Пост не знайдено")
    comments = (await session.execute(
        text("SELECT * FROM forum_comments WHERE post_id = :id ORDER BY created_at ASC"),
        {"id": post_id},
    )).mappings().all()
    return {**dict(row), "comments": [dict(c) for c in comments]}


@app.post("/forum/posts/{post_id}/comments", status_code=201)
async def create_forum_comment(
    post_id: str,
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    cid = str(uuid.uuid4())
    await session.execute(
        text("INSERT INTO forum_comments (id, post_id, user_id, content) VALUES (:id, :pid, :uid, :content)"),
        {"id": cid, "pid": post_id, "uid": uid, "content": body.get("content", "")},
    )
    await session.commit()
    return {"id": cid, "post_id": post_id, "user_id": uid,
            "content": body.get("content", ""), "created_at": datetime.utcnow().isoformat()}


# ============================================================================
# AI ANALYZE  (/api/analyze/*)
# ============================================================================

def _subscription_usage(plan: str, docs_used: int, docs_limit: int | None) -> dict:
    return {"docs_used": docs_used, "docs_limit": docs_limit}


async def _get_usage(session: AsyncSession, user_id: str) -> dict:
    row = (
        await session.execute(
            text("SELECT plan, docs_used, docs_limit FROM subscriptions WHERE user_id = :uid LIMIT 1"),
            {"uid": user_id},
        )
    ).mappings().first()
    if not row:
        return {"docs_used": 0, "docs_limit": 5}
    return _subscription_usage(row["plan"], row["docs_used"], row["docs_limit"])


async def _run_ai_intake(
    file_bytes: bytes,
    file_name: str,
    jurisdiction: str,
    mode: str,
    api_key: str | None,
) -> dict:
    """Call Anthropic API for deep document intake analysis."""
    if not api_key:
        return _demo_intake(file_name, jurisdiction)

    try:
        import anthropic  # noqa: PLC0415

        client = anthropic.AsyncAnthropic(api_key=api_key)
        text_preview = file_bytes.decode("utf-8", errors="replace")[:10000]
        max_tokens = 4096 if mode == "deep" else 2048

        prompt = f"""Ти — юридичний аналітик. Проаналізуй такий документ і поверни ТІЛЬКИ валідний JSON без markdown та додаткового тексту.

Документ (перші символи):
{text_preview}

Поверни JSON з такими полями:
{{
  "classified_type": "рядок (договір/позов/рішення/судовий наказ/інше)",
  "document_language": "uk або en або інше",
  "primary_party_role": "позивач/відповідач/кредитор/боржник/інше або null",
  "identified_parties": [{{"role": "...", "name": "..."}}],
  "subject_matter": "короткий опис предмету (1-2 речення)",
  "financial_exposure_amount": число або null,
  "financial_exposure_currency": "UAH/USD/EUR або null",
  "financial_exposure_type": "борг/штраф/компенсація або null",
  "document_date": "YYYY-MM-DD або null",
  "deadline_from_document": "YYYY-MM-DD або null",
  "urgency_level": "low/medium/high/critical",
  "risk_level_legal": "low/medium/high/critical",
  "risk_level_procedural": "low/medium/high/critical",
  "risk_level_financial": "low/medium/high/critical",
  "detected_issues": [
    {{"issue_type": "...", "severity": "low/medium/high/critical", "description": "...", "impact": "..."}}
  ],
  "classifier_confidence": число від 0 до 1,
  "tags": ["рядок", ...]
}}"""

        message = await client.messages.create(
            model=os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6"),
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = message.content[0].text.strip()
        # Strip markdown code fences if present
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except Exception as e:
        print(f"[analyze] AI error: {e}")
        return _demo_intake(file_name, jurisdiction)


def _demo_intake(file_name: str, jurisdiction: str) -> dict:
    return {
        "classified_type": "контракт",
        "document_language": "uk",
        "primary_party_role": None,
        "identified_parties": [],
        "subject_matter": f"[Demo] Файл: {file_name}. Налаштуйте ANTHROPIC_API_KEY для реального аналізу.",
        "financial_exposure_amount": None,
        "financial_exposure_currency": None,
        "financial_exposure_type": None,
        "document_date": None,
        "deadline_from_document": None,
        "urgency_level": "medium",
        "risk_level_legal": "medium",
        "risk_level_procedural": "low",
        "risk_level_financial": "low",
        "detected_issues": [
            {
                "issue_type": "demo",
                "severity": "low",
                "description": "Demo режим — API ключ не налаштовано",
                "impact": "Додайте ANTHROPIC_API_KEY у змінні оточення Railway",
            }
        ],
        "classifier_confidence": 0.0,
        "tags": ["demo"],
    }


def _intake_row_to_dict(row: Any, usage: dict) -> dict:
    r = dict(row)
    for field in ("identified_parties", "detected_issues", "tags"):
        if isinstance(r.get(field), str):
            try:
                r[field] = json.loads(r[field])
            except Exception:
                r[field] = []
    r["usage"] = usage
    r["cache_hit"] = False
    return r


def _build_intake_structured_json(ai_result: dict, file_name: str, jurisdiction: str) -> dict:
    issues = ai_result.get("detected_issues", []) or []
    return {
        "meta": {
            "file_name": file_name,
            "jurisdiction": jurisdiction,
            "classified_type": ai_result.get("classified_type", "unknown"),
            "language": ai_result.get("document_language"),
        },
        "parties": ai_result.get("identified_parties", []) or [],
        "timeline": {
            "document_date": ai_result.get("document_date"),
            "deadline_from_document": ai_result.get("deadline_from_document"),
            "urgency_level": ai_result.get("urgency_level", "medium"),
        },
        "risk_profile": {
            "legal": ai_result.get("risk_level_legal", "medium"),
            "procedural": ai_result.get("risk_level_procedural", "low"),
            "financial": ai_result.get("risk_level_financial", "low"),
            "issues_count": len(issues),
        },
        "subject_matter": ai_result.get("subject_matter"),
        "issues": issues,
        "tags": ai_result.get("tags", []) or [],
    }


def _extract_contract_pain_points(contract_text: str) -> dict:
    text_l = contract_text.lower()
    pain_points = []
    if "штраф" in text_l or "пеня" in text_l:
        pain_points.append({"type": "penalty_exposure", "severity": "high", "hint": "Високі штрафні санкції або пеня."})
    if "односторон" in text_l and ("розір" in text_l or "відмов" in text_l):
        pain_points.append({"type": "unilateral_termination", "severity": "high", "hint": "Одностороннє розірвання/відмова може бути ризикованим."})
    if "автоматичн" in text_l and "пролонгац" in text_l:
        pain_points.append({"type": "auto_renewal", "severity": "medium", "hint": "Автопролонгація без явного consent."})
    if "конфіденц" in text_l and "відповідальн" not in text_l:
        pain_points.append({"type": "nda_gap", "severity": "medium", "hint": "Є згадка про конфіденційність без чіткої відповідальності."})
    if "форс-мажор" not in text_l and "force majeure" not in text_l:
        pain_points.append({"type": "force_majeure_missing", "severity": "low", "hint": "Відсутній або слабкий блок форс-мажору."})
    return {
        "pain_points": pain_points,
        "risk_score": min(100, 20 * len(pain_points)),
        "total_flags": len(pain_points),
    }


@app.post("/api/analyze/intake")
async def analyze_intake(
    file: UploadFile = File(...),
    jurisdiction: str = Form("UA"),
    case_id: str = Form(None),
    mode: str = Query("standard"),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    t0 = time.monotonic()
    uid = str(current_user["id"])

    # Read file bytes (no disk write needed — ephemeral containers)
    try:
        file_bytes = await file.read()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Не вдалося прочитати файл: {e}")

    file_name = file.filename or "document"
    api_key = os.getenv("ANTHROPIC_API_KEY") or os.getenv("OPENAI_API_KEY")

    ai_result = await _run_ai_intake(file_bytes, file_name, jurisdiction, mode, api_key)

    elapsed_ms = int((time.monotonic() - t0) * 1000)
    rid = str(uuid.uuid4())

    await session.execute(
        text("""
            INSERT INTO document_intakes (
                id, user_id, case_id, source_file_name,
                classified_type, document_language, jurisdiction,
                primary_party_role, identified_parties, subject_matter,
                financial_exposure_amount, financial_exposure_currency, financial_exposure_type,
                document_date, deadline_from_document, urgency_level,
                risk_level_legal, risk_level_procedural, risk_level_financial,
                detected_issues, classifier_confidence, classifier_model,
                raw_text_preview, tags
            ) VALUES (
                :id, :uid, :case_id, :fname,
                :ctype, :lang, :jur,
                :role, :parties::jsonb, :subject,
                :amount, :currency, :exp_type,
                :doc_date, :deadline, :urgency,
                :rl_legal, :rl_proc, :rl_fin,
                :issues::jsonb, :conf, :model,
                :preview, :tags::jsonb
            )
        """),
        {
            "id": rid,
            "uid": uid,
            "case_id": case_id or None,
            "fname": file_name,
            "ctype": ai_result.get("classified_type", "unknown"),
            "lang": ai_result.get("document_language"),
            "jur": jurisdiction,
            "role": ai_result.get("primary_party_role"),
            "parties": json.dumps(ai_result.get("identified_parties", []), ensure_ascii=False),
            "subject": ai_result.get("subject_matter"),
            "amount": ai_result.get("financial_exposure_amount"),
            "currency": ai_result.get("financial_exposure_currency"),
            "exp_type": ai_result.get("financial_exposure_type"),
            "doc_date": ai_result.get("document_date"),
            "deadline": ai_result.get("deadline_from_document"),
            "urgency": ai_result.get("urgency_level", "medium"),
            "rl_legal": ai_result.get("risk_level_legal", "medium"),
            "rl_proc": ai_result.get("risk_level_procedural", "low"),
            "rl_fin": ai_result.get("risk_level_financial", "low"),
            "issues": json.dumps(ai_result.get("detected_issues", []), ensure_ascii=False),
            "conf": ai_result.get("classifier_confidence"),
            "model": os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6") if api_key else "demo",
            "preview": file_bytes.decode("utf-8", errors="replace")[:500],
            "tags": json.dumps(ai_result.get("tags", []), ensure_ascii=False),
        },
    )
    await session.commit()

    usage = await _get_usage(session, uid)
    row = (
        await session.execute(
            text("SELECT * FROM document_intakes WHERE id = :id LIMIT 1"), {"id": rid}
        )
    ).mappings().first()
    response = _intake_row_to_dict(row, usage)
    response["structured_output"] = _build_intake_structured_json(ai_result, file_name, jurisdiction)
    return response


@app.post("/api/analyze/intake-stream")
async def analyze_intake_stream(
    file: UploadFile = File(...),
    jurisdiction: str = Form("UA"),
    case_id: str = Form(None),
    mode: str = Query("standard"),
    current_user: dict = Depends(get_current_user),
):
    """SSE streaming wrapper.

    ВАЖЛИВО: не використовуємо Depends(get_session) тут, бо FastAPI закриває
    сесію після return StreamingResponse, ДО того як генератор почне yielding.
    Натомість будуємо відповідь з ai_result в пам'яті, а в БД зберігаємо через
    окрему сесію з AsyncSessionLocal.
    """
    file_bytes = await file.read()
    file_name = file.filename or "document"
    api_key = os.getenv("ANTHROPIC_API_KEY") or os.getenv("OPENAI_API_KEY")
    uid = str(current_user["id"])

    async def event_stream():
        yield "data: " + json.dumps({"event": "start", "message": "Аналіз розпочато…"}, ensure_ascii=False) + "\n\n"
        await asyncio.sleep(0)

        # AI аналіз (завжди повертає dict, ніколи не кидає виняток)
        try:
            ai_result = await _run_ai_intake(file_bytes, file_name, jurisdiction, mode, api_key)
        except Exception as exc:
            print(f"[intake-stream] AI error: {exc}")
            ai_result = _demo_intake(file_name, jurisdiction)

        rid = str(uuid.uuid4())
        now = datetime.utcnow().isoformat()
        model_name = os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6") if api_key else "demo"

        # Будуємо відповідь ОДРАЗУ з ai_result — без читання з БД
        result: dict = {
            "id": rid,
            "user_id": uid,
            "source_file_name": file_name,
            "classified_type": ai_result.get("classified_type", "unknown"),
            "document_language": ai_result.get("document_language"),
            "jurisdiction": jurisdiction,
            "primary_party_role": ai_result.get("primary_party_role"),
            "identified_parties": ai_result.get("identified_parties", []),
            "subject_matter": ai_result.get("subject_matter"),
            "financial_exposure_amount": ai_result.get("financial_exposure_amount"),
            "financial_exposure_currency": ai_result.get("financial_exposure_currency"),
            "financial_exposure_type": ai_result.get("financial_exposure_type"),
            "document_date": ai_result.get("document_date"),
            "deadline_from_document": ai_result.get("deadline_from_document"),
            "urgency_level": ai_result.get("urgency_level", "medium"),
            "risk_level_legal": ai_result.get("risk_level_legal", "medium"),
            "risk_level_procedural": ai_result.get("risk_level_procedural", "low"),
            "risk_level_financial": ai_result.get("risk_level_financial", "low"),
            "detected_issues": ai_result.get("detected_issues", []),
            "classifier_confidence": ai_result.get("classifier_confidence"),
            "classifier_model": model_name,
            "raw_text_preview": file_bytes.decode("utf-8", errors="replace")[:500],
            "tags": ai_result.get("tags", []),
            "created_at": now,
            "cache_hit": False,
            "usage": {"docs_used": 0, "docs_limit": 5},
        }
        result["structured_output"] = _build_intake_structured_json(ai_result, file_name, jurisdiction)

        # Зберігаємо в БД з власною сесією (best-effort, не блокує відповідь)
        try:
            from app.db import AsyncSessionLocal  # noqa: PLC0415
            async with AsyncSessionLocal() as db:
                await db.execute(
                    text("""
                        INSERT INTO document_intakes (
                            id, user_id, case_id, source_file_name,
                            classified_type, document_language, jurisdiction,
                            primary_party_role, identified_parties, subject_matter,
                            financial_exposure_amount, financial_exposure_currency, financial_exposure_type,
                            document_date, deadline_from_document, urgency_level,
                            risk_level_legal, risk_level_procedural, risk_level_financial,
                            detected_issues, classifier_confidence, classifier_model,
                            raw_text_preview, tags
                        ) VALUES (
                            :id, :uid, :case_id, :fname,
                            :ctype, :lang, :jur,
                            :role, :parties::jsonb, :subject,
                            :amount, :currency, :exp_type,
                            :doc_date, :deadline, :urgency,
                            :rl_legal, :rl_proc, :rl_fin,
                            :issues::jsonb, :conf, :model,
                            :preview, :tags::jsonb
                        )
                    """),
                    {
                        "id": rid, "uid": uid, "case_id": case_id or None, "fname": file_name,
                        "ctype": result["classified_type"], "lang": result["document_language"],
                        "jur": jurisdiction, "role": result["primary_party_role"],
                        "parties": json.dumps(result["identified_parties"], ensure_ascii=False),
                        "subject": result["subject_matter"],
                        "amount": result["financial_exposure_amount"],
                        "currency": result["financial_exposure_currency"],
                        "exp_type": result["financial_exposure_type"],
                        "doc_date": result["document_date"],
                        "deadline": result["deadline_from_document"],
                        "urgency": result["urgency_level"],
                        "rl_legal": result["risk_level_legal"],
                        "rl_proc": result["risk_level_procedural"],
                        "rl_fin": result["risk_level_financial"],
                        "issues": json.dumps(result["detected_issues"], ensure_ascii=False),
                        "conf": result["classifier_confidence"],
                        "model": model_name,
                        "preview": result["raw_text_preview"],
                        "tags": json.dumps(result["tags"], ensure_ascii=False),
                    },
                )
                await db.commit()
                # Оновлюємо usage з БД
                usage = await _get_usage(db, uid)
                result["usage"] = usage
        except Exception as db_exc:
            print(f"[intake-stream] DB save error: {db_exc}")

        yield "data: " + json.dumps({"event": "result", "result": result}, ensure_ascii=False) + "\n\n"
        yield "data: " + json.dumps({"event": "done"}) + "\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


@app.get("/api/analyze/history")
async def get_analyze_history(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    rows = (
        await session.execute(
            text("SELECT * FROM contract_analyses WHERE user_id = :uid ORDER BY created_at DESC LIMIT 50"),
            {"uid": uid},
        )
    ).mappings().all()
    usage = await _get_usage(session, uid)
    items = []
    for r in rows:
        d = dict(r)
        for f in ("critical_risks", "medium_risks", "ok_points", "recommendations"):
            if isinstance(d.get(f), str):
                try:
                    d[f] = json.loads(d[f])
                except Exception:
                    d[f] = []
        d["usage"] = usage
        items.append(d)
    return {"total": len(items), "items": items, "usage": usage}


@app.post("/api/analyze/process")
async def process_contract_analysis(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    t0 = time.monotonic()
    contract_text = body.get("contract_text", "")
    file_name = body.get("file_name", "document.txt")
    api_key = os.getenv("ANTHROPIC_API_KEY") or os.getenv("OPENAI_API_KEY")

    from app.services.intake_analyzer import run_intake_analysis  # noqa: PLC0415
    result = await run_intake_analysis(
        contract_text.encode("utf-8"),
        file_name,
        body.get("mode"),
    )
    heuristic = _extract_contract_pain_points(contract_text)
    elapsed_ms = int((time.monotonic() - t0) * 1000)
    rid = str(uuid.uuid4())
    await session.execute(
        text("""
            INSERT INTO contract_analyses (
                id, user_id, file_name, file_url, file_size,
                risk_level, critical_risks, medium_risks, ok_points,
                recommendations, summary, ai_model, tokens_used, processing_time_ms
            ) VALUES (
                :id, :uid, :fname, :furl, :fsize,
                :rl, :cr::jsonb, :mr::jsonb, :ok::jsonb,
                :rec::jsonb, :summary, :model, :tokens, :ms
            )
        """),
        {
            "id": rid, "uid": uid, "fname": file_name,
            "furl": body.get("file_url"), "fsize": body.get("file_size"),
            "rl": result.get("risk_level", "medium"),
            "cr": json.dumps(result.get("critical_risks", []), ensure_ascii=False),
            "mr": json.dumps(result.get("medium_risks", []), ensure_ascii=False),
            "ok": json.dumps(result.get("ok_points", []), ensure_ascii=False),
            "rec": json.dumps(result.get("recommendations", []), ensure_ascii=False),
            "summary": result.get("summary"),
            "model": result.get("ai_model", "demo"),
            "tokens": result.get("tokens_used", 0),
            "ms": elapsed_ms,
        },
    )
    await session.commit()
    usage = await _get_usage(session, uid)
    return {
        "id": rid,
        "user_id": uid,
        "file_name": file_name,
        "file_url": body.get("file_url"),
        "file_size": body.get("file_size"),
        "contract_type": None,
        "risk_level": result.get("risk_level", "medium"),
        "critical_risks": result.get("critical_risks", []),
        "medium_risks": result.get("medium_risks", []),
        "ok_points": result.get("ok_points", []),
        "recommendations": result.get("recommendations", []),
        "pain_points": heuristic["pain_points"],
        "risk_score": heuristic["risk_score"],
        "total_flags": heuristic["total_flags"],
        "summary": result.get("summary"),
        "ai_model": result.get("ai_model", "demo"),
        "tokens_used": result.get("tokens_used", 0),
        "processing_time_ms": elapsed_ms,
        "tags": [],
        "created_at": datetime.utcnow().isoformat(),
        "usage": usage,
    }


@app.get("/api/analyze/{analysis_id}")
async def get_analysis(
    analysis_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    usage = await _get_usage(session, uid)

    # Try document_intakes first
    row = (
        await session.execute(
            text("SELECT * FROM document_intakes WHERE id = :id AND user_id = :uid LIMIT 1"),
            {"id": analysis_id, "uid": uid},
        )
    ).mappings().first()
    if row:
        return _intake_row_to_dict(row, usage)

    # Fallback to contract_analyses
    row = (
        await session.execute(
            text("SELECT * FROM contract_analyses WHERE id = :id AND user_id = :uid LIMIT 1"),
            {"id": analysis_id, "uid": uid},
        )
    ).mappings().first()
    if row:
        d = dict(row)
        d["usage"] = usage
        return d

    raise HTTPException(status_code=404, detail="Аналіз не знайдено")


@app.patch("/api/analyze/{analysis_id}")
async def update_analysis(
    analysis_id: str,
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    allowed = {"classified_type", "subject_matter", "urgency_level", "risk_level_legal",
                "risk_level_procedural", "risk_level_financial", "tags"}
    fields = {k: v for k, v in body.items() if k in allowed and v is not None}
    if fields:
        sets = ", ".join(f"{k} = :{k}" for k in fields)
        if "tags" in fields:
            fields["tags"] = json.dumps(fields["tags"], ensure_ascii=False)
            sets = sets.replace("tags = :tags", "tags = :tags::jsonb")
        fields.update({"id": analysis_id, "uid": uid})
        await session.execute(
            text(f"UPDATE document_intakes SET {sets} WHERE id = :id AND user_id = :uid"),
            fields,
        )
        await session.commit()
    return {"status": "ok"}


@app.post("/api/analyze/gdpr-check")
async def gdpr_check(
    body: dict,
    current_user: dict = Depends(get_current_user),
):
    text_content = body.get("text", "")
    api_key = os.getenv("ANTHROPIC_API_KEY") or os.getenv("OPENAI_API_KEY")

    if not api_key:
        return {
            "report": "[Demo] GDPR-перевірка потребує ANTHROPIC_API_KEY",
            "compliant": True,
            "issues": [],
            "personal_data_found": [],
            "recommendations": ["Налаштуйте ANTHROPIC_API_KEY у Railway"],
        }

    try:
        import anthropic  # noqa: PLC0415
        client = anthropic.AsyncAnthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))
        prompt = (
            f"Проаналізуй текст на GDPR-відповідність і поверни JSON:\n"
            f"{{'compliant': bool, 'report': '...', 'issues': [...], "
            f"'personal_data_found': [{{'type': '...', 'count': N}}], "
            f"'recommendations': [...]}}\n\nТекст:\n{text_content[:5000]}"
        )
        msg = await client.messages.create(
            model=os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6"),
            max_tokens=1024,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except Exception as e:
        return {
            "report": f"Помилка аналізу: {e}",
            "compliant": True,
            "issues": [],
            "personal_data_found": [],
            "recommendations": [],
        }


@app.get("/api/analyze/{intake_id}/comments")
async def get_intake_comments(
    intake_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    rows = (
        await session.execute(
            text("""
                SELECT c.id, c.intake_id, c.user_id, u.full_name as user_name, c.content, c.created_at
                FROM intake_comments c
                LEFT JOIN users u ON u.id = c.user_id
                WHERE c.intake_id = :iid
                ORDER BY c.created_at ASC
            """),
            {"iid": intake_id},
        )
    ).mappings().all()
    return [dict(r) for r in rows]


@app.post("/api/analyze/{intake_id}/comments")
async def create_intake_comment(
    intake_id: str,
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    cid = str(uuid.uuid4())
    content = body.get("content", "").strip()
    if not content:
        raise HTTPException(status_code=400, detail="Порожній коментар")
    await session.execute(
        text("INSERT INTO intake_comments (id, intake_id, user_id, content) VALUES (:id, :iid, :uid, :content)"),
        {"id": cid, "iid": intake_id, "uid": uid, "content": content},
    )
    await session.commit()
    return {
        "id": cid,
        "intake_id": intake_id,
        "user_id": uid,
        "user_name": current_user.get("full_name"),
        "content": content,
        "created_at": datetime.utcnow().isoformat(),
    }


@app.delete("/api/analyze/{intake_id}/comments/{comment_id}")
async def delete_intake_comment(
    intake_id: str,
    comment_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    await session.execute(
        text("DELETE FROM intake_comments WHERE id = :cid AND intake_id = :iid AND user_id = :uid"),
        {"cid": comment_id, "iid": intake_id, "uid": uid},
    )
    await session.commit()
    return {"status": "ok"}


@app.post("/api/analyze/{intake_id}/precedent-map")
async def create_precedent_map(
    intake_id: str,
    limit: int = Query(10),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    intake = (
        await session.execute(
            text("SELECT subject_matter, tags FROM document_intakes WHERE id = :id AND user_id = :uid LIMIT 1"),
            {"id": intake_id, "uid": uid},
        )
    ).mappings().first()
    if not intake:
        raise HTTPException(status_code=404, detail="Intake не знайдено")

    tags = intake.get("tags")
    if isinstance(tags, str):
        try:
            tags = json.loads(tags)
        except Exception:
            tags = []
    terms = [str(t).strip() for t in (tags or []) if str(t).strip()]
    if intake.get("subject_matter"):
        terms.insert(0, str(intake["subject_matter"]))
    query_used = " ".join(terms[:4]).strip()
    if not query_used:
        query_used = "цивільний спір"

    rows = (
        await session.execute(
            text("""
                SELECT id, case_number, court_name, decision_date, summary, relevance_score
                FROM case_law_items
                WHERE user_id = :uid
                  AND (
                    summary ILIKE :q
                    OR case_number ILIKE :q
                    OR court_name ILIKE :q
                  )
                ORDER BY COALESCE(relevance_score, 0) DESC, decision_date DESC
                LIMIT :lim
            """),
            {"uid": uid, "q": f"%{query_used[:120]}%", "lim": max(1, min(limit, 30))},
        )
    ).mappings().all()

    groups_map: dict[str, list] = {}
    refs = []
    for row in rows:
        court = str(row.get("court_name") or "Невідомий суд")
        groups_map.setdefault(court, []).append(
            {
                "id": str(row["id"]),
                "case_number": row.get("case_number"),
                "decision_date": str(row.get("decision_date") or ""),
                "summary": row.get("summary"),
                "relevance_score": float(row.get("relevance_score") or 0),
            }
        )
        refs.append({"id": str(row["id"]), "court_name": court, "case_number": row.get("case_number")})

    groups = [{"court_name": court, "items": items} for court, items in groups_map.items()]
    return {
        "intake_id": intake_id,
        "query_used": query_used,
        "groups": groups,
        "refs": refs,
    }


# ============================================================================
# STRATEGY  (/api/strategy/*)
# ============================================================================

async def _ai_json(prompt: str, max_tokens: int = 4096) -> dict | list:
    """Call Anthropic and parse JSON from response."""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        return {}
    try:
        import anthropic  # noqa: PLC0415
        client = anthropic.AsyncAnthropic(api_key=api_key)
        msg = await client.messages.create(
            model=os.getenv("ANTHROPIC_MODEL", "claude-sonnet-4-6"),
            max_tokens=max_tokens,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = msg.content[0].text.strip()
        raw = re.sub(r"^```(?:json)?\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        return json.loads(raw)
    except Exception as e:
        print(f"[ai_json] error: {e}")
        return {}


@app.post("/api/strategy/blueprint")
async def create_strategy_blueprint(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    intake_id = body.get("intake_id", "")
    uid = str(current_user["id"])

    # Отримуємо intake дані якщо є
    intake_row = None
    if intake_id:
        intake_row = (await session.execute(
            text("SELECT * FROM document_intakes WHERE id = :id AND user_id = :uid LIMIT 1"),
            {"id": intake_id, "uid": uid},
        )).mappings().first()

    context = ""
    if intake_row:
        context = (
            f"Тип документу: {intake_row.get('classified_type')}\n"
            f"Предмет спору: {intake_row.get('subject_matter')}\n"
            f"Рівень ризику (правовий): {intake_row.get('risk_level_legal')}\n"
            f"Рівень ризику (процесуальний): {intake_row.get('risk_level_procedural')}\n"
            f"Рівень ризику (фінансовий): {intake_row.get('risk_level_financial')}\n"
            f"Виявлені проблеми: {intake_row.get('detected_issues')}\n"
        )

    prompt = f"""Ти — стратегічний юридичний аналітик. На основі аналізу документу розроби детальну правову стратегію.

{context}

Поверни ТІЛЬКИ валідний JSON без markdown:
{{
  "immediate_actions": [{{"priority": "high", "action": "...", "deadline": "...", "rationale": "..."}}],
  "procedural_roadmap": [{{"step": 1, "title": "...", "description": "...", "timeline": "..."}}],
  "evidence_strategy": [{{"type": "...", "description": "...", "importance": "high/medium/low"}}],
  "negotiation_playbook": [{{"scenario": "...", "approach": "...", "concessions": "...", "red_lines": "..."}}],
  "risk_heat_map": [{{"risk": "...", "probability": "high/medium/low", "impact": "high/medium/low", "mitigation": "..."}}],
  "critical_deadlines": [{{"event": "...", "date": null, "consequence": "..."}}],
  "swot_analysis": {{
    "strengths": ["..."], "weaknesses": ["..."],
    "opportunities": ["..."], "threats": ["..."]
  }},
  "win_probability": 0.65
}}"""

    blueprint = await _ai_json(prompt)
    if not blueprint:
        blueprint = {
            "immediate_actions": [{"priority": "high", "action": "Налаштуйте ANTHROPIC_API_KEY для реальної стратегії", "deadline": "зараз", "rationale": "Demo режим"}],
            "procedural_roadmap": [], "evidence_strategy": [], "negotiation_playbook": [],
            "risk_heat_map": [], "critical_deadlines": [],
            "swot_analysis": {"strengths": [], "weaknesses": [], "opportunities": [], "threats": []},
            "win_probability": None,
        }

    bid = str(uuid.uuid4())
    return {
        "id": bid, "intake_id": intake_id, "precedent_group_id": None,
        **blueprint,
        "financial_strategy": None, "created_at": datetime.utcnow().isoformat(),
    }


@app.post("/api/strategy/blueprint-stream")
async def strategy_blueprint_stream(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """SSE streaming для генерації стратегії."""
    intake_id = body.get("intake_id", "")
    uid = str(current_user["id"])

    async def stream():
        yield "data: " + json.dumps({"event": "start", "message": "Генерую стратегію…"}) + "\n\n"
        # Re-use the non-streaming endpoint logic
        try:
            # Get intake context
            context = ""
            if intake_id:
                row = (await session.execute(
                    text("SELECT * FROM document_intakes WHERE id = :id LIMIT 1"),
                    {"id": intake_id},
                )).mappings().first()
                if row:
                    context = f"Тип: {row.get('classified_type')}, Предмет: {row.get('subject_matter')}"

            yield "data: " + json.dumps({"event": "progress", "message": "Аналізую документ…"}) + "\n\n"
            await asyncio.sleep(0)

            prompt = f"""Стратегічний юридичний аналітик. Контекст: {context}
Поверни JSON стратегії з полями: immediate_actions, procedural_roadmap, evidence_strategy,
negotiation_playbook, risk_heat_map, critical_deadlines, swot_analysis, win_probability."""
            blueprint = await _ai_json(prompt)
            if not blueprint:
                blueprint = {"immediate_actions": [], "procedural_roadmap": [], "evidence_strategy": [],
                             "negotiation_playbook": [], "risk_heat_map": [], "critical_deadlines": [],
                             "swot_analysis": {"strengths": [], "weaknesses": [], "opportunities": [], "threats": []},
                             "win_probability": None}

            result = {"id": str(uuid.uuid4()), "intake_id": intake_id, **blueprint,
                      "created_at": datetime.utcnow().isoformat()}
            yield "data: " + json.dumps({"event": "result", "result": result}, ensure_ascii=False) + "\n\n"
        except Exception as e:
            yield "data: " + json.dumps({"event": "error", "message": str(e)}) + "\n\n"
        yield "data: " + json.dumps({"event": "done"}) + "\n\n"

    return StreamingResponse(stream(), media_type="text/event-stream")


@app.post("/api/strategy/simulate-judge")
async def simulate_judge(
    body: dict,
    current_user: dict = Depends(get_current_user),
):
    intake_id = body.get("intake_id", "")
    prompt = f"""Симулюй роль судді. Інтейк ID: {intake_id}.
Поверни JSON: {{"judge_persona": "...", "key_vulnerabilities": ["..."], "strong_points": ["..."],
"procedural_risks": ["..."], "suggested_corrections": ["..."],
"judge_commentary": "...", "decision_rationale": "..."}}"""
    result = await _ai_json(prompt, max_tokens=2048)
    if not result:
        result = {
            "judge_persona": "Суддя загальної юрисдикції",
            "key_vulnerabilities": ["Налаштуйте ANTHROPIC_API_KEY для реальної симуляції"],
            "strong_points": [], "procedural_risks": [], "suggested_corrections": [],
            "judge_commentary": "Demo режим", "decision_rationale": "Demo режим",
        }
    return {"id": str(uuid.uuid4()), "intake_id": intake_id, **result,
            "created_at": datetime.utcnow().isoformat()}


@app.post("/api/generate-with-strategy")
async def generate_with_strategy(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    doc_type = body.get("doc_type", "позов")
    blueprint_id = body.get("blueprint_id")
    intake_id = body.get("intake_id")
    form_data = body.get("form_data", {})

    context_parts = [f"Тип документу: {doc_type}"]
    if form_data:
        context_parts.append(f"Дані форми: {json.dumps(form_data, ensure_ascii=False)[:1000]}")

    prompt = (
        f"Склади юридичний документ типу '{doc_type}' для України.\n"
        f"Контекст: {', '.join(context_parts)}\n\n"
        "Поверни повний текст документу українською мовою у форматі JSON: "
        "{\"text\": \"повний текст документу\", \"title\": \"назва\"}"
    )
    result = await _ai_json(prompt, max_tokens=4096)
    generated_text = result.get("text", "[Demo] Налаштуйте ANTHROPIC_API_KEY") if result else "[Demo]"
    title = result.get("title", f"{doc_type} — {datetime.utcnow().strftime('%d.%m.%Y')}")

    did = str(uuid.uuid4())
    try:
        await session.execute(
            text("""
                INSERT INTO generated_documents (id, user_id, document_type, document_category,
                    title, generated_text, used_ai, ai_model)
                VALUES (:id, :uid, :dtype, 'civil', :title, :text, true, :model)
            """),
            {"id": did, "uid": uid, "dtype": doc_type, "title": title,
             "text": generated_text, "model": os.getenv("ANTHROPIC_MODEL", "demo")},
        )
        await session.commit()
    except Exception as e:
        print(f"[generate-with-strategy] DB error: {e}")

    return {"id": did, "user_id": uid, "document_type": doc_type, "title": title,
            "generated_text": generated_text, "ai_model": os.getenv("ANTHROPIC_MODEL", "demo"),
            "created_at": datetime.utcnow().isoformat()}


# ============================================================================
# DOCUMENTS (missing endpoints)
# ============================================================================

@app.post("/api/documents/generate-stream")
async def generate_document_stream(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    """SSE streaming для генерації документів."""
    uid = str(current_user["id"])
    doc_type = body.get("doc_type", "позов")
    form_data = body.get("form_data", {})

    async def stream():
        yield "data: " + json.dumps({"event": "start", "message": "Генерую документ…"}) + "\n\n"
        await asyncio.sleep(0)
        try:
            prompt = (
                f"Склади юридичний документ типу '{doc_type}' для України.\n"
                f"Дані: {json.dumps(form_data, ensure_ascii=False)[:2000]}\n"
                "Поверни JSON: {\"text\": \"...\", \"title\": \"...\"}"
            )
            result = await _ai_json(prompt, max_tokens=4096)
            generated_text = result.get("text", "[Demo]") if result else "[Demo]"
            title = result.get("title", doc_type) if result else doc_type
            did = str(uuid.uuid4())
            try:
                from app.db import AsyncSessionLocal  # noqa: PLC0415
                async with AsyncSessionLocal() as db:
                    await db.execute(text("""
                        INSERT INTO generated_documents
                            (id, user_id, document_type, document_category, title, generated_text, used_ai, ai_model)
                        VALUES (:id, :uid, :dtype, 'civil', :title, :text, true, :model)
                    """), {"id": did, "uid": uid, "dtype": doc_type, "title": title,
                           "text": generated_text, "model": os.getenv("ANTHROPIC_MODEL", "demo")})
                    await db.commit()
            except Exception as dbe:
                print(f"[generate-stream] DB error: {dbe}")
            doc = {"id": did, "user_id": uid, "document_type": doc_type, "title": title,
                   "generated_text": generated_text, "ai_model": os.getenv("ANTHROPIC_MODEL", "demo"),
                   "created_at": datetime.utcnow().isoformat()}
            yield "data: " + json.dumps({"event": "result", "result": doc}, ensure_ascii=False) + "\n\n"
        except Exception as e:
            yield "data: " + json.dumps({"event": "error", "message": str(e)}) + "\n\n"
        yield "data: " + json.dumps({"event": "done"}) + "\n\n"

    return StreamingResponse(stream(), media_type="text/event-stream")


@app.post("/api/documents/bulk-delete")
async def bulk_delete_documents(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    ids = body.get("ids", [])
    if ids:
        for doc_id in ids:
            await session.execute(
                text("DELETE FROM generated_documents WHERE id = :id AND user_id = :uid"),
                {"id": doc_id, "uid": uid},
            )
        await session.commit()
    return {"deleted": len(ids)}


@app.post("/api/documents/processual-gate-check")
async def processual_gate_check(body: dict, current_user: dict = Depends(get_current_user)):
    return {"passed": True, "blockers": [], "warnings": []}


@app.post("/api/documents/bulk-processual-repair")
async def bulk_processual_repair(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    ids = [str(x) for x in (body.get("ids") or []) if str(x).strip()]
    if not ids:
        return {
            "status": "ok",
            "requested": 0,
            "processed": 0,
            "repaired": 0,
            "missing_ids": [],
            "items": [],
        }

    rows = (
        await session.execute(
            text(
                """
                SELECT id, document_type, generated_text
                FROM generated_documents
                WHERE user_id = :uid AND id = ANY(:ids)
                """
            ),
            {"uid": uid, "ids": ids},
        )
    ).mappings().all()
    rows_by_id = {str(r["id"]): r for r in rows}
    missing_ids = [doc_id for doc_id in ids if doc_id not in rows_by_id]

    repaired_count = 0
    items: list[dict[str, Any]] = []

    for doc_id in ids:
        row = rows_by_id.get(doc_id)
        if not row:
            continue
        original_text = row.get("generated_text") or ""
        repaired_text, blockers_before, blockers_after = await _repair_processual_text(
            row["document_type"], original_text
        )
        repaired = repaired_text != original_text
        if repaired:
            repaired_count += 1
            await session.execute(
                text(
                    """
                    UPDATE generated_documents
                    SET generated_text = :text,
                        preview_text = :preview,
                        updated_at = NOW()
                    WHERE id = :id AND user_id = :uid
                    """
                ),
                {"id": doc_id, "uid": uid, "text": repaired_text, "preview": repaired_text[:200]},
            )
        items.append(
            {
                "id": doc_id,
                "status": "repaired" if repaired else "checked",
                "repaired": repaired,
                "is_valid": len(blockers_after) == 0,
                "blockers": blockers_after,
                "blockers_before": blockers_before,
            }
        )

    await session.commit()
    return {
        "status": "ok",
        "requested": len(ids),
        "processed": len(rows),
        "repaired": repaired_count,
        "missing_ids": missing_ids,
        "items": items,
    }


# ============================================================================
# E-COURT  (/api/e-court/*)
# ============================================================================

@app.get("/api/e-court/courts")
async def get_courts(
    region: str | None = Query(None),
    current_user: dict = Depends(get_current_user),
):
    courts = [
        {"id": "kyiv_app", "name": "Київський апеляційний суд", "type": "appeal", "region": "Київ"},
        {"id": "kyiv_com", "name": "Господарський суд м. Київ", "type": "commercial", "region": "Київ"},
        {"id": "kyiv_adm", "name": "Окружний адміністративний суд м. Київ", "type": "administrative", "region": "Київ"},
        {"id": "supreme", "name": "Верховний Суд", "type": "supreme", "region": "Київ"},
        {"id": "lviv_app", "name": "Львівський апеляційний суд", "type": "appeal", "region": "Львів"},
        {"id": "kharkiv_app", "name": "Харківський апеляційний суд", "type": "appeal", "region": "Харків"},
        {"id": "odessa_app", "name": "Одеський апеляційний суд", "type": "appeal", "region": "Одеса"},
        {"id": "dnipro_app", "name": "Дніпровський апеляційний суд", "type": "appeal", "region": "Дніпро"},
    ]
    if region:
        courts = [c for c in courts if region.lower() in c["region"].lower()]
    return {"items": courts}


@app.get("/api/e-court/hearings")
async def get_ecourt_hearings(
    case_id: str | None = Query(None),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    params: dict[str, Any] = {"uid": uid}
    filters = ["s.user_id = :uid", "s.status IN ('submitted', 'accepted', 'in_review')"]
    if case_id:
        filters.append("d.case_id = :case_id")
        params["case_id"] = case_id
    where = " AND ".join(filters)
    rows = (
        await session.execute(
            text(
                f"""
                SELECT
                    s.id,
                    s.status,
                    s.court_name,
                    s.submitted_at AS hearing_at,
                    COALESCE(d.title, 'Подання до е-суду') AS title,
                    d.case_id
                FROM ecourt_submissions s
                LEFT JOIN generated_documents d ON d.id = s.document_id
                WHERE {where}
                ORDER BY s.submitted_at DESC
                LIMIT 100
                """
            ),
            params,
        )
    ).mappings().all()
    items = [
        {
            "id": str(r["id"]),
            "case_id": str(r["case_id"]) if r.get("case_id") else None,
            "title": r.get("title") or "Подання до е-суду",
            "court_name": r.get("court_name") or "",
            "status": r.get("status") or "submitted",
            "hearing_at": r["hearing_at"].isoformat() if hasattr(r.get("hearing_at"), "isoformat") else r.get("hearing_at"),
            "source": "ecourt_submissions",
        }
        for r in rows
    ]
    return {"items": items, "total": len(items)}


# ============================================================================
# KNOWLEDGE BASE  (/api/knowledge-base/*)
# ============================================================================

@app.post("/api/knowledge-base/")
async def create_knowledge_entry(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    kid = str(uuid.uuid4())
    title = body.get("title", "")
    content = body.get("content", "")
    category = body.get("category")
    try:
        await session.execute(
            text("""
                INSERT INTO knowledge_entries (id, user_id, title, content, category)
                VALUES (:id, :uid, :title, :content, :category)
            """),
            {"id": kid, "uid": uid, "title": title, "content": content, "category": category},
        )
        await session.commit()
    except Exception as e:
        print(f"[knowledge-base] DB error: {e}")
        return {"id": kid, "user_id": uid, "title": title, "content": content,
                "category": category, "created_at": datetime.utcnow().isoformat()}
    return {"id": kid, "user_id": uid, "title": title, "content": content,
            "category": category, "created_at": datetime.utcnow().isoformat()}


@app.delete("/api/knowledge-base/{entry_id}")
async def delete_knowledge_entry(
    entry_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    try:
        await session.execute(
            text("DELETE FROM knowledge_entries WHERE id = :id AND user_id = :uid"),
            {"id": entry_id, "uid": uid},
        )
        await session.commit()
    except Exception:
        pass
    return {"status": "ok"}


# ============================================================================
# TEAM / AUTH
# ============================================================================

@app.get("/api/auth/team/users")
async def get_auth_team_users(current_user: dict = Depends(get_current_user)):
    return {"items": [{"id": str(current_user["id"]), "email": current_user["email"],
                       "role": current_user.get("role", "user"), "full_name": current_user.get("full_name")}]}


@app.patch("/api/auth/team/users/role")
async def update_team_user_role(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    user_id = body.get("user_id")
    role = body.get("role", "user")
    if user_id:
        try:
            await session.execute(
                text("UPDATE users SET role = :role WHERE id = :uid"),
                {"role": role, "uid": user_id},
            )
            await session.commit()
        except Exception:
            pass
    return {"status": "ok"}


# ============================================================================
# USER PREFERENCES
# ============================================================================

@app.get("/api/users/me/preferences")
async def get_preferences(current_user: dict = Depends(get_current_user)):
    return {"theme": "dark", "language": "uk", "notifications": True}


@app.patch("/api/users/me/preferences")
async def update_preferences(body: dict, current_user: dict = Depends(get_current_user)):
    return {"status": "ok", **body}


# ============================================================================
# KNOWLEDGE ENTRIES (повний CRUD для /api/knowledge-base)
# ============================================================================

@app.get("/api/knowledge-base/")
async def get_knowledge_entries_full(
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    try:
        rows = (await session.execute(
            text("SELECT * FROM knowledge_entries WHERE user_id = :uid ORDER BY created_at DESC LIMIT 100"),
            {"uid": uid},
        )).mappings().all()
        return {"items": [dict(r) for r in rows]}
    except Exception:
        return {"items": []}


# ── OpenDataBot proxy ─────────────────────────────────────────────────────────

_ODB_KEY = os.getenv("OPENDATABOT_API_KEY", "")
_ODB_BASE = "https://api.opendatabot.ua"


async def _odb_get(path: str, params: dict | None = None) -> Any:
    """Proxy a GET request to OpenDataBot API."""
    if not _ODB_KEY:
        raise HTTPException(status_code=503, detail="OPENDATABOT_API_KEY не налаштований")
    headers = {"apikey": _ODB_KEY}
    async with httpx.AsyncClient(timeout=30) as client:
        resp = await client.get(f"{_ODB_BASE}{path}", params=params, headers=headers)
    if resp.status_code == 404:
        raise HTTPException(status_code=404, detail="Справу не знайдено в OpenDataBot")
    if resp.status_code == 402:
        raise HTTPException(status_code=402, detail="Ліміт запитів OpenDataBot вичерпано")
    if not resp.is_success:
        raise HTTPException(status_code=resp.status_code, detail=f"OpenDataBot error: {resp.text[:200]}")
    return resp.json()


def _extract_pdf_links_from_case_payload(payload: dict) -> list[str]:
    links: list[str] = []

    def collect(value: Any) -> None:
        if isinstance(value, dict):
            for v in value.values():
                collect(v)
            return
        if isinstance(value, list):
            for item in value:
                collect(item)
            return
        if isinstance(value, str) and value.lower().startswith(("http://", "https://")) and ".pdf" in value.lower():
            links.append(value)

    collect(payload)
    dedup: list[str] = []
    seen: set[str] = set()
    for link in links:
        if link not in seen:
            seen.add(link)
            dedup.append(link)
    return dedup


@app.get("/api/opendatabot/court-cases/{case_number}")
async def get_court_case(
    case_number: str,
    judgment_code: int | None = Query(None),
    current_user: dict = Depends(get_current_user),
):
    params: dict = {"number": case_number}
    if judgment_code:
        params["judgment_code"] = judgment_code
    data = await _odb_get("/v1/court/case", params=params)
    return data


@app.get("/api/opendatabot/usage")
async def get_opendatabot_usage(current_user: dict = Depends(get_current_user)):
    if not _ODB_KEY:
        return {"limit": 0, "used": 0, "remaining": 0, "expires_at": None, "api_url": _ODB_BASE}
    data = await _odb_get("/v1/account/usage")
    return {
        "limit": data.get("limit", 0),
        "used": data.get("used", 0),
        "remaining": data.get("remaining", 0),
        "expires_at": data.get("expires_at"),
        "api_url": _ODB_BASE,
    }


@app.get("/api/opendatabot/company/{code}")
async def get_company(code: str, current_user: dict = Depends(get_current_user)):
    data = await _odb_get(f"/v1/company/{code}")
    return data


# ============================================================================
# E-COURT — повний CRUD з таблицею ecourt_submissions
# ============================================================================

@app.get("/api/e-court/history")
async def get_ecourt_history(
    page: int = Query(1, ge=1),
    page_size: int = Query(20),
    status: str | None = Query(None),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    params: dict = {"uid": uid}
    conditions = ["user_id = :uid"]
    if status:
        conditions.append("status = :status")
        params["status"] = status
    where = " AND ".join(conditions)
    count_row = (await session.execute(
        text(f"SELECT COUNT(*) FROM ecourt_submissions WHERE {where}"), params
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text(f"SELECT * FROM ecourt_submissions WHERE {where} ORDER BY submitted_at DESC LIMIT :lim OFFSET :off"),
        {**params, "lim": page_size, "off": offset},
    )).mappings().all()
    return {"total": count_row, "page": page, "page_size": page_size,
            "pages": max(1, -(-int(count_row) // page_size)), "items": [dict(r) for r in rows]}


@app.post("/api/e-court/submit")
async def submit_to_ecourt(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    sid = str(uuid.uuid4())
    ext_id = f"ES-{sid[:8].upper()}"
    await session.execute(
        text("""
            INSERT INTO ecourt_submissions
                (id, user_id, document_id, provider, external_submission_id,
                 status, court_name, signer_method, note)
            VALUES (:id, :uid, :doc_id, 'manual', :ext_id,
                    'submitted', :court, :signer, :note)
        """),
        {
            "id": sid, "uid": uid,
            "doc_id": body.get("document_id"),
            "ext_id": ext_id,
            "court": body.get("court_name", ""),
            "signer": body.get("signer_method", "manual"),
            "note": body.get("note"),
        },
    )
    await session.commit()
    submission = {
        "id": sid, "document_id": body.get("document_id"), "provider": "manual",
        "external_submission_id": ext_id, "status": "submitted",
        "court_name": body.get("court_name", ""), "signer_method": body.get("signer_method"),
        "tracking_url": None, "error_message": None,
        "submitted_at": datetime.utcnow().isoformat(), "updated_at": datetime.utcnow().isoformat(),
    }
    await _audit_log(session, uid, "ecourt_submit", "ecourt_submission", sid)
    return {"status": "submitted", "submission": submission}


@app.get("/api/e-court/{submission_id}/status")
async def get_ecourt_status(
    submission_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    row = (await session.execute(
        text("SELECT * FROM ecourt_submissions WHERE id = :id LIMIT 1"), {"id": submission_id}
    )).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Подання не знайдено")
    return {"submission": dict(row)}


@app.post("/api/e-court/{submission_id}/sync-status")
async def sync_ecourt_status(
    submission_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    row = (await session.execute(
        text("SELECT * FROM ecourt_submissions WHERE id = :id AND user_id = :uid LIMIT 1"),
        {"id": submission_id, "uid": uid},
    )).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Подання не знайдено")
    submission = dict(row)
    synced_live, sync_reason = await _sync_submission_with_adapter(submission)

    status_value = submission.get("status") or "submitted"
    if synced_live and status_value == "submitted":
        status_value = "in_review"

    await session.execute(
        text(
            """
            UPDATE ecourt_submissions
            SET status = :status, updated_at = NOW()
            WHERE id = :id AND user_id = :uid
            """
        ),
        {"status": status_value, "id": submission_id, "uid": uid},
    )
    await session.commit()

    refreshed = (await session.execute(
        text("SELECT * FROM ecourt_submissions WHERE id = :id AND user_id = :uid LIMIT 1"),
        {"id": submission_id, "uid": uid},
    )).mappings().first()
    return {"submission": dict(refreshed) if refreshed else submission, "synced_live": synced_live, "sync_reason": sync_reason}


@app.get("/api/e-court/public-search")
async def ecourt_public_search(
    case_number: str = Query(""),
    current_user: dict = Depends(get_current_user),
):
    """Пошук справи через OpenDataBot якщо є ключ, інакше — порожній результат."""
    if not case_number.strip():
        raise HTTPException(status_code=400, detail="Номер справи обов'язковий")
    if _ODB_KEY:
        try:
            data = await _odb_get("/v1/court/case", params={"number": case_number})
            return {
                "status": "found", "case_number": case_number,
                "assignments": data.get("assignments", []),
                "history": data.get("history", []),
            }
        except HTTPException as e:
            if e.status_code == 404:
                return {"status": "not_found", "case_number": case_number, "assignments": [], "history": []}
            raise
    return {"status": "no_key", "case_number": case_number, "assignments": [], "history": []}


@app.get("/api/e-court/public-search/pdf")
async def ecourt_public_search_pdf(
    case_number: str = Query(""),
    current_user: dict = Depends(get_current_user),
):
    if not case_number.strip():
        raise HTTPException(status_code=400, detail="Номер справи обов'язковий")
    if not _ODB_KEY:
        return {
            "status": "no_key",
            "case_number": case_number,
            "pdf_links": [],
            "message": "OPENDATABOT_API_KEY не налаштований",
        }
    payload = await _odb_get("/v1/court/case", params={"number": case_number})
    pdf_links = _extract_pdf_links_from_case_payload(payload if isinstance(payload, dict) else {})
    return {"status": "ok", "case_number": case_number, "pdf_links": pdf_links}


# ============================================================================
# MONITORING — повний CRUD watch-items з OpenDataBot перевіркою
# ============================================================================

@app.get("/api/monitoring/watch-items")
async def get_watch_items_list(
    page: int = Query(1, ge=1),
    page_size: int = Query(20),
    registry_type: str | None = Query(None),
    status: str | None = Query(None),
    query: str | None = Query(None),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    params: dict = {"uid": uid}
    conditions = ["user_id = :uid"]
    if registry_type:
        conditions.append("registry_type = :rtype")
        params["rtype"] = registry_type
    if status:
        conditions.append("status = :status")
        params["status"] = status
    if query:
        conditions.append("(entity_name ILIKE :q OR identifier ILIKE :q)")
        params["q"] = f"%{query}%"
    where = " AND ".join(conditions)
    count_row = (await session.execute(
        text(f"SELECT COUNT(*) FROM registry_watch_items WHERE {where}"), params
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text(f"SELECT * FROM registry_watch_items WHERE {where} ORDER BY created_at DESC LIMIT :lim OFFSET :off"),
        {**params, "lim": page_size, "off": offset},
    )).mappings().all()
    return {"total": count_row, "page": page, "page_size": page_size,
            "pages": max(1, -(-int(count_row) // page_size)), "items": [dict(r) for r in rows]}


@app.post("/api/monitoring/watch-items", status_code=201)
async def create_watch_item_full(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    wid = str(uuid.uuid4())
    from datetime import timedelta  # noqa: PLC0415
    next_check = datetime.utcnow() + timedelta(hours=int(body.get("check_interval_hours", 24)))
    await session.execute(
        text("""
            INSERT INTO registry_watch_items
                (id, user_id, source, registry_type, identifier, entity_name,
                 status, check_interval_hours, next_check_at, notes)
            VALUES (:id, :uid, :source, :rtype, :ident, :name,
                    'active', :interval, :next_check, :notes)
        """),
        {
            "id": wid, "uid": uid,
            "source": body.get("source", "opendatabot"),
            "rtype": body.get("registry_type", "company"),
            "ident": body.get("identifier", ""),
            "name": body.get("entity_name", ""),
            "interval": body.get("check_interval_hours", 24),
            "next_check": next_check,
            "notes": body.get("notes"),
        },
    )
    await session.commit()
    row = (await session.execute(
        text("SELECT * FROM registry_watch_items WHERE id = :id LIMIT 1"), {"id": wid}
    )).mappings().first()
    item = dict(row) if row else {"id": wid, "user_id": uid, **body}
    return {"status": "created", "item": item}


@app.delete("/api/monitoring/watch-items/{item_id}")
async def delete_watch_item(
    item_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    await session.execute(
        text("DELETE FROM registry_watch_items WHERE id = :id AND user_id = :uid"),
        {"id": item_id, "uid": uid},
    )
    await session.commit()
    return {"status": "deleted", "id": item_id}


@app.post("/api/monitoring/watch-items/{item_id}/check")
async def check_watch_item(
    item_id: str,
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    row = (await session.execute(
        text("SELECT * FROM registry_watch_items WHERE id = :id AND user_id = :uid LIMIT 1"),
        {"id": item_id, "uid": uid},
    )).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Запис спостереження не знайдено")

    item = dict(row)
    event_type = "state_unchanged"
    severity = "info"
    title = f"Перевірено: {item.get('entity_name', item.get('identifier'))}"
    details: dict = {}

    # Якщо є OPENDATABOT_API_KEY — реальна перевірка
    if _ODB_KEY and item.get("registry_type") == "company":
        try:
            data = await _odb_get(f"/v1/company/{item['identifier']}")
            new_status = data.get("status", "")
            prev_snapshot = item.get("latest_snapshot") or {}
            if isinstance(prev_snapshot, str):
                try:
                    prev_snapshot = json.loads(prev_snapshot)
                except Exception:
                    prev_snapshot = {}
            if new_status != prev_snapshot.get("status"):
                event_type = "state_changed"
                severity = "warning"
                title = f"Зміна статусу: {item.get('entity_name')}"
            details = {"old_status": prev_snapshot.get("status"), "new_status": new_status}
            await session.execute(
                text("""UPDATE registry_watch_items SET last_checked_at = NOW(),
                    latest_snapshot = :snap::jsonb, last_change_at = CASE WHEN :changed THEN NOW() ELSE last_change_at END
                    WHERE id = :id"""),
                {"snap": json.dumps(data, ensure_ascii=False), "changed": event_type == "state_changed", "id": item_id},
            )
            await session.execute(
                text("""
                    INSERT INTO registry_snapshots (id, user_id, watch_item_id, snapshot, source)
                    VALUES (:id, :uid, :wid, :snap::jsonb, 'opendatabot')
                """),
                {"id": str(uuid.uuid4()), "uid": uid, "wid": item_id, "snap": json.dumps(data, ensure_ascii=False)},
            )
        except Exception as e:
            details = {"error": str(e)}
    else:
        await session.execute(
            text("UPDATE registry_watch_items SET last_checked_at = NOW() WHERE id = :id"),
            {"id": item_id},
        )

    eid = str(uuid.uuid4())
    await session.execute(
        text("""
            INSERT INTO registry_events (id, user_id, watch_item_id, event_type, severity, title, details)
            VALUES (:id, :uid, :wid, :etype, :sev, :title, :details::jsonb)
        """),
        {"id": eid, "uid": uid, "wid": item_id, "etype": event_type,
         "sev": severity, "title": title, "details": json.dumps(details, ensure_ascii=False)},
    )
    await session.commit()
    row = (await session.execute(
        text("SELECT * FROM registry_watch_items WHERE id = :id LIMIT 1"), {"id": item_id}
    )).mappings().first()
    return {"status": "checked", "item": dict(row) if row else item,
            "event_id": eid, "event_type": event_type}


@app.get("/api/monitoring/watch-items/{item_id}/history")
async def get_watch_item_history(
    item_id: str,
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text("""
            SELECT id, watch_item_id, snapshot, source, created_at
            FROM registry_snapshots
            WHERE user_id = :uid AND watch_item_id = :wid
            ORDER BY created_at DESC
            LIMIT :lim OFFSET :off
        """),
        {"uid": uid, "wid": item_id, "lim": page_size, "off": offset},
    )).mappings().all()
    total = (await session.execute(
        text("SELECT COUNT(*) FROM registry_snapshots WHERE user_id = :uid AND watch_item_id = :wid"),
        {"uid": uid, "wid": item_id},
    )).scalar() or 0
    items = []
    for row in rows:
        d = dict(row)
        if isinstance(d.get("snapshot"), str):
            try:
                d["snapshot"] = json.loads(d["snapshot"])
            except Exception:
                d["snapshot"] = {}
        items.append(d)
    return {"total": int(total), "page": page, "page_size": page_size, "items": items}


@app.get("/api/monitoring/events")
async def get_monitoring_events(
    page: int = Query(1, ge=1),
    page_size: int = Query(20),
    watch_item_id: str | None = Query(None),
    severity: str | None = Query(None),
    event_type: str | None = Query(None),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    params: dict = {"uid": uid}
    conditions = ["user_id = :uid"]
    if watch_item_id:
        conditions.append("watch_item_id = :wid")
        params["wid"] = watch_item_id
    if severity:
        conditions.append("severity = :sev")
        params["sev"] = severity
    if event_type:
        conditions.append("event_type = :etype")
        params["etype"] = event_type
    where = " AND ".join(conditions)
    count_row = (await session.execute(
        text(f"SELECT COUNT(*) FROM registry_events WHERE {where}"), params
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text(f"SELECT * FROM registry_events WHERE {where} ORDER BY created_at DESC LIMIT :lim OFFSET :off"),
        {**params, "lim": page_size, "off": offset},
    )).mappings().all()
    items = []
    for r in rows:
        d = dict(r)
        if isinstance(d.get("details"), str):
            try:
                d["details"] = json.loads(d["details"])
            except Exception:
                d["details"] = {}
        items.append(d)
    return {"total": count_row, "page": page, "page_size": page_size,
            "pages": max(1, -(-int(count_row) // page_size)), "items": items}


@app.post("/api/monitoring/check-due")
async def monitoring_check_due(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    limit = body.get("limit", 10)
    due_rows = (await session.execute(
        text("""
            SELECT * FROM registry_watch_items
            WHERE user_id = :uid AND status = 'active'
              AND (next_check_at IS NULL OR next_check_at <= NOW())
            ORDER BY next_check_at ASC NULLS FIRST
            LIMIT :lim
        """),
        {"uid": uid, "lim": limit},
    )).mappings().all()

    checked = 0
    state_changed = 0
    for item in due_rows:
        item_id = str(item["id"])
        event_type = "state_unchanged"
        if _ODB_KEY and item.get("registry_type") == "company":
            try:
                data = await _odb_get(f"/v1/company/{item['identifier']}")
                prev = item.get("latest_snapshot") or {}
                if isinstance(prev, str):
                    try:
                        prev = json.loads(prev)
                    except Exception:
                        prev = {}
                if data.get("status") != prev.get("status"):
                    event_type = "state_changed"
                    state_changed += 1
                await session.execute(
                    text("UPDATE registry_watch_items SET last_checked_at=NOW(), latest_snapshot=:snap::jsonb, next_check_at=NOW() + (check_interval_hours || ' hours')::interval WHERE id=:id"),
                    {"snap": json.dumps(data, ensure_ascii=False), "id": item_id},
                )
            except Exception:
                pass
        else:
            await session.execute(
                text("UPDATE registry_watch_items SET last_checked_at=NOW(), next_check_at=NOW() + (check_interval_hours || ' hours')::interval WHERE id=:id"),
                {"id": item_id},
            )
        await session.execute(
            text("INSERT INTO registry_events (id,user_id,watch_item_id,event_type,severity,title,details) VALUES (:id,:uid,:wid,:etype,'info',:title,'{}')"),
            {"id": str(uuid.uuid4()), "uid": uid, "wid": item_id,
             "etype": event_type, "title": f"Авто-перевірка: {item.get('entity_name', item['identifier'])}"},
        )
        checked += 1
    await session.commit()
    return {"status": "done", "scanned": len(due_rows), "checked": checked, "state_changed": state_changed}


# ============================================================================
# DOCUMENT VERSIONS — таблиця document_versions
# ============================================================================

@app.get("/api/documents/{doc_id}/versions")
async def get_document_versions(
    doc_id: str,
    page: int = Query(1, ge=1),
    page_size: int = Query(20),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    # Verify ownership
    doc = (await session.execute(
        text("SELECT id FROM generated_documents WHERE id = :id AND user_id = :uid LIMIT 1"),
        {"id": doc_id, "uid": uid},
    )).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Документ не знайдено")
    count_row = (await session.execute(
        text("SELECT COUNT(*) FROM document_versions WHERE document_id = :did"), {"did": doc_id}
    )).scalar() or 0
    offset = (page - 1) * page_size
    rows = (await session.execute(
        text("SELECT id, document_id, version_number, action, created_at FROM document_versions WHERE document_id = :did ORDER BY version_number DESC LIMIT :lim OFFSET :off"),
        {"did": doc_id, "lim": page_size, "off": offset},
    )).mappings().all()
    return {"document_id": doc_id, "total": count_row, "page": page, "page_size": page_size,
            "pages": max(1, -(-int(count_row) // page_size)), "items": [dict(r) for r in rows]}


@app.get("/api/documents/{doc_id}/versions/{version_id}")
async def get_document_version_detail(
    doc_id: str,
    version_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    row = (await session.execute(
        text("SELECT * FROM document_versions WHERE id = :vid AND document_id = :did LIMIT 1"),
        {"vid": version_id, "did": doc_id},
    )).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Версія не знайдена")
    return dict(row)


@app.get("/api/documents/{doc_id}/versions/{version_id}/diff")
async def get_document_version_diff(
    doc_id: str,
    version_id: str,
    against: str = Query("current"),
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    target = (await session.execute(
        text("SELECT * FROM document_versions WHERE id = :vid AND document_id = :did LIMIT 1"),
        {"vid": version_id, "did": doc_id},
    )).mappings().first()
    if not target:
        raise HTTPException(status_code=404, detail="Версія не знайдена")

    if against == "current":
        current_doc = (await session.execute(
            text("SELECT generated_text FROM generated_documents WHERE id = :id LIMIT 1"),
            {"id": doc_id},
        )).first()
        against_text = current_doc[0] if current_doc else ""
        against_ver_num = None
    else:
        other = (await session.execute(
            text("SELECT * FROM document_versions WHERE id = :vid AND document_id = :did LIMIT 1"),
            {"vid": against, "did": doc_id},
        )).mappings().first()
        against_text = dict(other).get("generated_text", "") if other else ""
        against_ver_num = dict(other).get("version_number") if other else None

    target_lines = (target.get("generated_text") or "").splitlines()
    against_lines = (against_text or "").splitlines()
    added = sum(1 for l in target_lines if l not in against_lines)
    removed = sum(1 for l in against_lines if l not in target_lines)
    return {
        "document_id": doc_id, "target_version_id": version_id,
        "target_version_number": target.get("version_number"),
        "against": against, "against_version_number": against_ver_num,
        "diff_text": f"+{added} рядків, -{removed} рядків",
        "added_lines": added, "removed_lines": removed,
    }


@app.post("/api/documents/{doc_id}/versions/{version_id}/restore")
async def restore_document_version(
    doc_id: str,
    version_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    version_row = (await session.execute(
        text("SELECT * FROM document_versions WHERE id = :vid AND document_id = :did LIMIT 1"),
        {"vid": version_id, "did": doc_id},
    )).mappings().first()
    if not version_row:
        raise HTTPException(status_code=404, detail="Версія не знайдена")

    v = dict(version_row)
    # Save current as new version before restoring
    cur = (await session.execute(
        text("SELECT generated_text FROM generated_documents WHERE id = :id AND user_id = :uid LIMIT 1"),
        {"id": doc_id, "uid": uid},
    )).first()
    if not cur:
        raise HTTPException(status_code=404, detail="Документ не знайдено")

    max_ver = (await session.execute(
        text("SELECT COALESCE(MAX(version_number), 0) FROM document_versions WHERE document_id = :did"),
        {"did": doc_id},
    )).scalar() or 0
    new_ver = max_ver + 1
    backup_id = str(uuid.uuid4())
    await session.execute(
        text("INSERT INTO document_versions (id, document_id, version_number, action, generated_text) VALUES (:id, :did, :ver, 'auto_backup', :text)"),
        {"id": backup_id, "did": doc_id, "ver": new_ver, "text": cur[0]},
    )
    restore_ver = new_ver + 1
    await session.execute(
        text("UPDATE generated_documents SET generated_text = :text WHERE id = :id AND user_id = :uid"),
        {"text": v.get("generated_text", ""), "id": doc_id, "uid": uid},
    )
    restore_id = str(uuid.uuid4())
    await session.execute(
        text("INSERT INTO document_versions (id, document_id, version_number, action, generated_text) VALUES (:id, :did, :ver, 'restore', :text)"),
        {"id": restore_id, "did": doc_id, "ver": restore_ver, "text": v.get("generated_text", "")},
    )
    await session.commit()
    await _audit_log(session, uid, "document_restore", "document", doc_id, {"from_version": version_id})
    return {
        "status": "restored", "id": doc_id,
        "restored_from_version_id": version_id,
        "restored_to_version_number": restore_ver,
        "has_docx_export": False, "has_pdf_export": False,
    }


# ============================================================================
# CASE-LAW SYNC — підтягуємо з OpenDataBot в case_law_items
# ============================================================================

@app.post("/api/case-law/sync")
async def sync_case_law(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    query = body.get("query", "")
    limit = min(int(body.get("limit", 5)), 10)
    job_id = str(uuid.uuid4())
    created = 0
    updated = 0

    if _ODB_KEY and query:
        try:
            data = await _odb_get("/v1/court/case", params={"number": query})
            cases = data if isinstance(data, list) else [data]
            for c in cases[:limit]:
                decision_id = str(c.get("case_id", c.get("id", str(uuid.uuid4()))))
                existing = (await session.execute(
                    text("SELECT id FROM case_law_items WHERE decision_id = :did AND user_id = :uid LIMIT 1"),
                    {"did": decision_id, "uid": uid},
                )).first()
                if existing:
                    await session.execute(
                        text("UPDATE case_law_items SET summary = :s, tags = :tags::jsonb WHERE decision_id = :did AND user_id = :uid"),
                        {"s": str(c.get("subject", ""))[:500], "tags": "[]", "did": decision_id, "uid": uid},
                    )
                    updated += 1
                else:
                    await session.execute(
                        text("""
                            INSERT INTO case_law_items
                                (id, user_id, source, decision_id, case_number, court_name, summary)
                            VALUES (:id, :uid, 'opendatabot', :did, :cnum, :court, :summary)
                        """),
                        {
                            "id": str(uuid.uuid4()), "uid": uid, "did": decision_id,
                            "cnum": str(c.get("case_number", query)),
                            "court": str(c.get("court_name", "")),
                            "summary": str(c.get("subject", ""))[:500],
                        },
                    )
                    created += 1
            await session.commit()
        except Exception as e:
            print(f"[case-law/sync] error: {e}")

    return {
        "status": "done", "created": created, "updated": updated,
        "total": created + updated, "sources": ["opendatabot"],
        "seed_fallback_used": not bool(_ODB_KEY), "fetched_counts": {"opendatabot": created + updated},
    }


@app.post("/api/case-law/import")
async def import_case_law(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    items = body.get("items", [])
    created = 0
    errors = []
    for item in items[:50]:
        try:
            await session.execute(
                text("""
                    INSERT INTO case_law_items
                        (id, user_id, source, decision_id, case_number, court_name,
                         judge_name, decision_date, doc_type, summary, full_text, tags)
                    VALUES
                        (:id, :uid, :src, :did, :cnum, :court,
                         :judge, :date, :dtype, :summary, :full_text, :tags::jsonb)
                    ON CONFLICT DO NOTHING
                """),
                {
                    "id": str(uuid.uuid4()), "uid": uid,
                    "src": item.get("source", "manual"),
                    "did": item.get("decision_id", str(uuid.uuid4())),
                    "cnum": item.get("case_number"),
                    "court": item.get("court_name"),
                    "judge": item.get("judge_name"),
                    "date": item.get("decision_date"),
                    "dtype": item.get("doc_type"),
                    "summary": item.get("summary", "")[:1000],
                    "full_text": item.get("full_text", ""),
                    "tags": json.dumps(item.get("tags", []), ensure_ascii=False),
                },
            )
            created += 1
        except Exception as e:
            errors.append(str(e))
    await session.commit()
    return {"imported": created, "skipped": len(items) - created, "errors": errors[:5]}


@app.post("/api/case-law/digest/generate")
async def generate_case_law_digest(
    body: dict,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    query = body.get("query", "")
    # Pull recent items from case_law_items
    rows = (await session.execute(
        text("SELECT case_number, court_name, summary FROM case_law_items WHERE user_id = :uid ORDER BY created_at DESC LIMIT 20"),
        {"uid": uid},
    )).mappings().all()
    items_text = "\n".join([f"- {r.get('case_number')}: {r.get('summary', '')}" for r in rows])

    api_key = os.getenv("ANTHROPIC_API_KEY")
    if api_key and rows:
        prompt = f"Склади короткий юридичний дайджест на основі цих справ:\n{items_text}\nЗапит: {query}\nПоверни JSON: {{\"title\": \"...\", \"summary\": \"...\"}}"
        result = await _ai_json(prompt, max_tokens=1024)
        title = result.get("title", f"Дайджест — {datetime.utcnow().strftime('%d.%m.%Y')}") if result else f"Дайджест"
        summary = result.get("summary", items_text[:500]) if result else items_text[:500]
    else:
        title = f"Дайджест — {datetime.utcnow().strftime('%d.%m.%Y')}"
        summary = items_text[:500] or "Немає даних у базі судової практики."

    did = str(uuid.uuid4())
    await session.execute(
        text("INSERT INTO case_law_digest (id, user_id, title, summary, query, items_count) VALUES (:id, :uid, :title, :summary, :query, :count)"),
        {"id": did, "uid": uid, "title": title, "summary": summary, "query": query, "count": len(rows)},
    )
    await session.commit()
    return {"id": did, "user_id": uid, "title": title, "summary": summary,
            "source": "ai" if api_key else "manual", "query": query,
            "items_count": len(rows), "created_at": datetime.utcnow().isoformat()}


# ============================================================================
# DOCUMENT EXPORT (DOCX)
# ============================================================================

@app.get("/api/documents/{doc_id}/export/docx")
async def export_document_docx(
    doc_id: str,
    current_user: dict = Depends(get_current_user),
    session: AsyncSession = Depends(get_session),
):
    uid = str(current_user["id"])
    row = (await session.execute(
        text("SELECT * FROM generated_documents WHERE id = :id AND user_id = :uid LIMIT 1"),
        {"id": doc_id, "uid": uid},
    )).mappings().first()
    if not row:
        raise HTTPException(status_code=404, detail="Документ не знайдено")
    doc = dict(row)

    try:
        from docx import Document as DocxDocument  # noqa: PLC0415
        from io import BytesIO  # noqa: PLC0415
        from fastapi.responses import Response  # noqa: PLC0415

        d = DocxDocument()
        d.add_heading(doc.get("title") or doc.get("document_type", "Документ"), 0)
        text_content = doc.get("generated_text", "")
        for paragraph in text_content.split("\n"):
            if paragraph.strip():
                d.add_paragraph(paragraph)
        buf = BytesIO()
        d.save(buf)
        buf.seek(0)
        filename = f"{doc.get('title', 'document')}.docx".replace("/", "_")
        return Response(
            content=buf.read(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=\"{filename}\""},
        )
    except ImportError:
        raise HTTPException(status_code=501, detail="python-docx не встановлено. Додайте python-docx до requirements.txt")


if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "main:app",
        host="0.0.0.0",
        port=int(os.getenv("PORT", "8000")),
        reload=os.getenv("RELOAD", "false").lower() == "true",
    )
