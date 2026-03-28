"""
Wave 4 · STORY-7 — Deterministic DocumentIR renderer.

Converts a DocumentIR into a properly formatted DOCX (and optionally PDF).
This module is NOT embedded in the generate endpoint — it is called only
after the IR has passed the final_render_gate.

DOCX generation uses python-docx (add `python-docx` to requirements.txt).
PDF generation converts the DOCX via LibreOffice subprocess (headless).
Both paths are gated by ENABLE_IR_RENDERER per doc_type (feature_flags.py).

Entry points:

    from .renderer import DocumentRenderer
    renderer = DocumentRenderer()
    docx_bytes = renderer.render_docx(ir)   # → bytes
    pdf_bytes  = renderer.render_pdf(ir)    # → bytes (via DOCX→PDF)

Legacy fallback:
    When ENABLE_IR_RENDERER is false for a doc_type, the existing
    generated_text blob is used as-is — NO breaking change.

Snapshot tests:  tests/test_renderer_snapshots.py
"""

from __future__ import annotations

import io
import os
import pathlib
import subprocess
import tempfile
from typing import TYPE_CHECKING

from fastapi import HTTPException

from .document_ir import DocumentIR, IRDocumentStatus
from .error_codes import RENDER_FAIL
from .feature_flags import flags

if TYPE_CHECKING:
    pass

# ---------------------------------------------------------------------------
# Optional python-docx import
# ---------------------------------------------------------------------------

try:
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt
    _DOCX_AVAILABLE = True
except ImportError:  # pragma: no cover
    _DOCX_AVAILABLE = False


# ---------------------------------------------------------------------------
# Renderer
# ---------------------------------------------------------------------------

class DocumentRenderer:
    """Convert a DocumentIR into DOCX / PDF bytes.

    Usage:
        renderer = DocumentRenderer()
        if flags.ir_renderer(ir.document_type):
            docx_bytes = renderer.render_docx(ir)
    """

    def render_docx(self, ir: DocumentIR) -> bytes:
        """Render DocumentIR → DOCX bytes.

        Raises HTTPException(500, RENDER_FAIL) on any rendering error.
        """
        if not _DOCX_AVAILABLE:
            raise HTTPException(
                status_code=500,
                detail={
                    "error_code": RENDER_FAIL,
                    "message": "python-docx не встановлено. Додайте 'python-docx' до requirements.txt.",
                },
            )
        try:
            doc = self._build_docx(ir)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail={
                    "error_code": RENDER_FAIL,
                    "message": f"Помилка формування DOCX: {exc}",
                },
            ) from exc

    def render_pdf(self, ir: DocumentIR) -> bytes:
        """Render DocumentIR → PDF bytes via LibreOffice headless conversion.

        Requires LibreOffice installed on the server:
            apt-get install -y libreoffice

        Raises HTTPException(500, RENDER_FAIL) if LibreOffice is unavailable.
        """
        docx_bytes = self.render_docx(ir)
        try:
            return _convert_docx_to_pdf(docx_bytes)
        except Exception as exc:
            raise HTTPException(
                status_code=500,
                detail={
                    "error_code": RENDER_FAIL,
                    "message": f"Помилка конвертації DOCX→PDF: {exc}",
                },
            ) from exc

    # ------------------------------------------------------------------
    # Internal DOCX construction
    # ------------------------------------------------------------------

    def _build_docx(self, ir: DocumentIR) -> "DocxDocument":
        doc = DocxDocument()

        # ── Page margins (A4, Ukrainian court standard) ────────────────
        for section in doc.sections:
            section.top_margin    = Cm(2.0)
            section.bottom_margin = Cm(2.0)
            section.left_margin   = Cm(3.0)   # court left margin
            section.right_margin  = Cm(1.5)

        # ── 1. Court header block ──────────────────────────────────────
        if ir.header.court_name:
            _add_paragraph(doc, ir.header.court_name,
                           bold=True, align="right", size=12)

        # ── 2. Parties block ──────────────────────────────────────────
        if ir.parties:
            doc.add_paragraph()  # blank line
            for party in ir.parties:
                line = f"{party.role.upper()}: {party.name}"
                if party.address:
                    line += f",  {party.address}"
                if party.identifier:
                    line += f"  (ЄДРПОУ/ІПН: {party.identifier})"
                _add_paragraph(doc, line, size=11)

        # ── 3. Title ──────────────────────────────────────────────────
        doc.add_paragraph()
        _add_paragraph(doc, ir.header.title.upper(),
                       bold=True, align="center", size=14)
        doc.add_paragraph()

        # ── 4. Facts ──────────────────────────────────────────────────
        if ir.facts:
            _add_paragraph(doc, "ОБСТАВИНИ СПРАВИ", bold=True, size=12)
            for i, fact in enumerate(ir.facts, 1):
                _add_paragraph(doc, f"{i}. {fact.text}", size=11, indent=Cm(1.25))
            doc.add_paragraph()

        # ── 5. Legal basis ────────────────────────────────────────────
        if ir.legal_basis:
            _add_paragraph(doc, "ПРАВОВА ПІДСТАВА", bold=True, size=12)
            for thesis in ir.legal_basis:
                _add_paragraph(doc, thesis.text, size=11, indent=Cm(1.25))
                for cit_id in thesis.citations:
                    citation = next((c for c in ir.citations if c.id == cit_id), None)
                    if citation:
                        _add_paragraph(
                            doc,
                            f"[{citation.source_locator}]: «{citation.evidence_span[:120]}…»",
                            italic=True, size=10, indent=Cm(2.0),
                        )
            doc.add_paragraph()

        # ── 6. Claims (прохальна частина) ─────────────────────────────
        if ir.claims:
            _add_paragraph(doc, "ПРОШУ:", bold=True, size=12)
            for i, claim in enumerate(ir.claims, 1):
                text = f"{i}. {claim.text}"
                if claim.amount is not None:
                    text += f" ({claim.amount:,.2f} {claim.currency or 'UAH'})"
                _add_paragraph(doc, text, size=11, indent=Cm(1.25))
            doc.add_paragraph()

        # ── 7. Attachments ────────────────────────────────────────────
        if ir.attachments:
            _add_paragraph(doc, "ДОДАТКИ:", bold=True, size=11)
            for i, att in enumerate(ir.attachments, 1):
                marker = "[*]" if att.required else "[ ]"
                _add_paragraph(doc, f"{i}. {marker} {att.title}", size=10, indent=Cm(1.25))
            doc.add_paragraph()

        # ── 8. Signature block ────────────────────────────────────────
        if ir.signature_block:
            sig = ir.signature_block
            doc.add_paragraph()
            signer = sig.signer_name or "_______________________"
            role   = sig.signer_role or ""
            if role:
                _add_paragraph(doc, role, size=11, align="right")
            sig_line = f"{signer}  ________________________"
            _add_paragraph(doc, sig_line, size=11, align="right")
            if sig.date_placeholder:
                _add_paragraph(doc, "«___» ____________ 20__ р.", size=11, align="right")

        return doc


# ---------------------------------------------------------------------------
# PDF conversion via LibreOffice
# ---------------------------------------------------------------------------

def _convert_docx_to_pdf(docx_bytes: bytes) -> bytes:
    with tempfile.TemporaryDirectory() as tmpdir:
        in_path = pathlib.Path(tmpdir) / "document.docx"
        in_path.write_bytes(docx_bytes)

        result = subprocess.run(
            [
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", tmpdir, str(in_path),
            ],
            capture_output=True, timeout=60,
        )
        if result.returncode != 0:
            raise RuntimeError(
                f"LibreOffice exited {result.returncode}: "
                f"{result.stderr.decode()[:200]}"
            )

        out_path = pathlib.Path(tmpdir) / "document.pdf"
        if not out_path.exists():
            raise RuntimeError("LibreOffice produced no PDF output.")
        return out_path.read_bytes()


# ---------------------------------------------------------------------------
# docx helper
# ---------------------------------------------------------------------------

def _add_paragraph(
    doc: "DocxDocument",
    text: str,
    bold: bool = False,
    italic: bool = False,
    align: str = "left",
    size: int = 12,
    indent: "Cm | None" = None,
) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent is not None:
        p.paragraph_format.left_indent = indent
